import os
import docx
from resume_customizer.detectors.project_detector import ProjectDetector
from resume_customizer.processors.document_processor import DocumentProcessor
from resume_customizer.processors.resume_processor import ResumeProcessor

def test_project_detection_fix():
    """Test that project detection correctly identifies company sections and job titles."""
    # Create a simple test document
    doc = docx.Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("ABC Company | Jan 2020 - Present")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("• Developed and maintained web applications using React and Node.js")
    doc.add_paragraph("• Implemented CI/CD pipelines for automated testing and deployment")
    doc.add_paragraph("XYZ Corporation | Mar 2018 - Dec 2019")
    doc.add_paragraph("Software Developer")
    doc.add_paragraph("• Built RESTful APIs using Python and Flask")
    doc.add_paragraph("• Collaborated with cross-functional teams to deliver features")
    
    # Initialize project detector
    detector = ProjectDetector()
    
    # Find projects in the resume
    projects = detector.find_projects(doc)
    
    # Verify correct number of projects detected
    assert len(projects) == 2, f"Expected 2 projects, but found {len(projects)}"
    
    # Verify first project details
    assert "ABC Company" in projects[0].name, f"First project name should contain 'ABC Company', got {projects[0].name}"
    assert projects[0].role == "Senior Software Engineer", f"First project role should be 'Senior Software Engineer', got {projects[0].role}"
    
    # Verify second project details
    assert "XYZ Corporation" in projects[1].name, f"Second project name should contain 'XYZ Corporation', got {projects[1].name}"
    assert projects[1].role == "Software Developer", f"Second project role should be 'Software Developer', got {projects[1].role}"
    
    print("✅ Project detection test passed!")

def test_point_placement_fix():
    """Test that points are added after the last bullet point in a project section."""
    # Create a simple test document
    doc = docx.Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("ABC Company | Jan 2020 - Present")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("• Developed and maintained web applications using React and Node.js")
    doc.add_paragraph("• Implemented CI/CD pipelines for automated testing and deployment")
    doc.add_paragraph("XYZ Corporation | Mar 2018 - Dec 2019")
    doc.add_paragraph("Software Developer")
    doc.add_paragraph("• Built RESTful APIs using Python and Flask")
    doc.add_paragraph("• Collaborated with cross-functional teams to deliver features")
    
    # Save the test document
    test_file = "test_document.docx"
    doc.save(test_file)
    
    try:
        # Initialize document processor
        processor = DocumentProcessor()
        
        # Define projects (normally these would come from ProjectDetector)
        projects = [
            {
                "name": "ABC Company | Jan 2020 - Present\nSenior Software Engineer",
                "start_index": 1,
                "end_index": 4,
                "bullet_indices": [3, 4]
            },
            {
                "name": "XYZ Corporation | Mar 2018 - Dec 2019\nSoftware Developer",
                "start_index": 5,
                "end_index": 8,
                "bullet_indices": [7, 8]
            }
        ]
        
        # Add points to the first project
        new_points = ["• Implemented new React components for improved user experience"]
        processor._add_points_to_project(doc, projects[0], new_points)
        
        # Save the modified document
        modified_file = "test_document_modified.docx"
        doc.save(modified_file)
        
        # Load the modified document to verify changes
        modified_doc = docx.Document(modified_file)
        paragraphs = [p.text for p in modified_doc.paragraphs]
        
        # The new point should be after the last bullet point of the first project
        assert paragraphs[5] == "• Implemented new React components for improved user experience", \
            f"New point should be after last bullet, got {paragraphs[5]}"
        
        # The second project should still start at the same position (but shifted by 1)
        assert "XYZ Corporation" in paragraphs[6], \
            f"Second project should start after new point, got {paragraphs[6]}"
        
        print("✅ Point placement test passed!")
        
    finally:
        # Clean up test files
        if os.path.exists(test_file):
            os.remove(test_file)
        if os.path.exists(modified_file):
            os.remove(modified_file)

def test_full_resume_processing():
    """Test the full resume processing flow with the fixes."""
    # This test requires actual resume files
    # Check if test resume files exist
    test_files = ["Resume Format 1.docx", "Resume Format 3.docx"]
    available_files = [f for f in test_files if os.path.exists(f)]
    
    if not available_files:
        print("⚠️ No test resume files found. Skipping full resume processing test.")
        return
    
    test_file = available_files[0]
    print(f"Testing with resume file: {test_file}")
    
    # Initialize resume processor
    processor = ResumeProcessor()
    
    # Process the resume with some test points
    test_points = [
        "Implemented React components for improved user experience",
        "Utilized Node.js for backend API development",
        "Deployed applications using Docker and Kubernetes"
    ]
    
    # Generate preview
    preview_file = f"preview_{test_file}"
    processor.generate_preview(test_file, preview_file, test_points)
    
    print(f"✅ Full resume processing test completed. Preview generated at {preview_file}")
    print("Please manually verify that points are correctly placed in the preview file.")

if __name__ == "__main__":
    print("Running tests for project detection and point placement fixes...")
    test_project_detection_fix()
    test_point_placement_fix()
    test_full_resume_processing()
    print("All tests completed!")