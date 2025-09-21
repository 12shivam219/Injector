"""
Resume processing module for Resume Customizer application.
Coordinates all resume processing operations with improved architecture and error handling.
"""

import threading
import functools
import queue
import os
import time
import psutil
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Callable, Tuple, Union
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field
import traceback

# External imports
from docx import Document
import streamlit as st

# Local imports
from infrastructure.monitoring.audit_logger import audit_logger
from infrastructure.utilities.logger import get_logger
import logging
from ..parsers.text_parser import parse_input_text, LegacyParser, get_parser
from ..parsers.restricted_text_parser import parse_input_text_restricted, RestrictedFormatError
from .document_processor import get_document_processor, FileProcessor
from ..email.email_handler import get_email_manager
from infrastructure.error_handling.file_processor import with_file_error_handling, FileProcessingError, FileErrorHandler
from infrastructure.error_handling.base import ErrorContext, ErrorSeverity
from core.errors import ProcessingError

class ResumeManager:
    """Manages resume processing operations."""
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        if not hasattr(self, '_initialized'):
            self._initialized = True
            # Initialize logger for this instance
            try:
                self.logger = get_logger()
            except Exception:
                self.logger = logging.getLogger(__name__)
                self.logger.setLevel(logging.INFO)
            self.document_processor = get_document_processor()
            self.email_manager = get_email_manager()
            self.progress_tracker = get_progress_tracker()
            self.error_recovery = get_error_recovery_manager()

    def process_resume(self, file_content: BytesIO, tech_stack: str) -> BytesIO:
        """Process a single resume file."""
        try:
            # Track progress
            if self.progress_tracker:
                self.progress_tracker.start_task(
                    "resume_processing",
                    "Processing resume",
                    total=100
                )

            # Parse tech stack
            parsed_points = parse_input_text_restricted(tech_stack)
            
            # Process document
            processed_doc = self.document_processor.process_document(
                file_content,
                parsed_points
            )

            # Record metrics
            record_metrics('resume_processed', 1, {'status': 'success'})

            return processed_doc

        except Exception as e:
            record_metrics('resume_processed', 1, {'status': 'failed'})
            raise

    def send_email(self, to_email: str, subject: str, body: str, attachment: BytesIO):
        """Send processed resume via email."""
        return self.email_manager.send_email(to_email, subject, body, attachment)

_resume_manager_instance = None
_resume_manager_lock = threading.Lock()

def get_resume_manager(version: str = "v2.0") -> ResumeManager:
    """Get the singleton instance of ResumeManager."""
    global _resume_manager_instance
    
    if _resume_manager_instance is None:
        with _resume_manager_lock:
            if _resume_manager_instance is None:
                _resume_manager_instance = ResumeManager()
    
    return _resume_manager_instance
try:
    from config import ERROR_MESSAGES, SUCCESS_MESSAGES, APP_CONFIG
except ImportError:
    ERROR_MESSAGES = {}
    SUCCESS_MESSAGES = {}
    APP_CONFIG = {}
    
try:
    from infrastructure.security.security_enhancements import rate_limit
except ImportError:
    def rate_limit(func): return func  # No-op decorator
    
from infrastructure.utilities.structured_logger import get_structured_logger, with_structured_logging
try:
    from infrastructure.utilities.structured_logger import processing_logger
except ImportError:
    processing_logger = get_structured_logger("processing")
    
from infrastructure.error_handling import ErrorHandler, ErrorContext, ErrorSeverity
# Removed circuit breaker import - monitoring disabled

try:
    from infrastructure.error_handling.recovery import ErrorRecoveryManager, get_error_recovery_manager
except ImportError:
    class ErrorRecoveryManager:
        def __init__(self): pass
    def get_error_recovery_manager(): return None
    
try:
    from infrastructure.monitoring.analytics import get_analytics_manager
    from infrastructure.monitoring.metrics import MetricType
    def record_metrics(category, **metric_kwargs):
        def decorator(func):
            @functools.wraps(func)
            def wrapper(*args, **kwargs):
                analytics = get_analytics_manager()
                start_time = time.time()
                try:
                    result = func(*args, **kwargs)
                    if analytics:
                        analytics.record_metric(
                            category,
                            'success',
                            1,
                            tags={'status': 'success', **metric_kwargs}
                        )
                    return result
                except Exception as e:
                    if analytics:
                        analytics.record_metric(
                            category,
                            'error',
                            1,
                            tags={'status': 'error', 'error_type': type(e).__name__, **metric_kwargs}
                        )
                    raise
                finally:
                    if analytics:
                        duration = time.time() - start_time
                        analytics.record_metric(
                            f"{category}.duration",
                            None,
                            duration,
                            metric_type=MetricType.TIMER,
                            tags=metric_kwargs
                        )
            return wrapper
        return decorator
except ImportError:
    def record_metrics(category, **metric_kwargs):
        def decorator(func):
            @functools.wraps(func)
            def wrapper(*args, **kwargs):
                return func(*args, **kwargs)
            return wrapper
        return decorator
    
try:
    from infrastructure.monitoring.progress import get_progress_tracker
except ImportError:
    def get_progress_tracker(): 
        class ProgressTracker:
            def track_progress(self, *args, **kwargs): pass
        return ProgressTracker()

# Initialize logger with fallback
try:
    logger = get_logger()
except Exception:
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)

# Constants
CACHE_TTL_SECONDS = 300  # 5 minutes
MAX_CACHE_SIZE = 100
DEFAULT_PROCESSING_TIMEOUT = 30


@dataclass
class ProcessingResult:
    """Data class for processing results with comprehensive information."""
    success: bool
    filename: str
    error: Optional[str] = None
    points_added: int = 0
    processing_time: float = 0.0
    tech_stacks_used: List[str] = field(default_factory=list)
    modified_content: Optional[bytes] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class CacheEntry:
    """Cache entry with TTL support."""
    result: ProcessingResult
    timestamp: datetime
    access_count: int = 0
    
    def is_expired(self) -> bool:
        """Check if cache entry is expired."""
        return datetime.now() - self.timestamp > timedelta(seconds=CACHE_TTL_SECONDS)


def profile_step(step_name: str):
    """Decorator to profile and log slow operations."""
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            start = time.time()
            try:
                result = func(*args, **kwargs)
                duration = time.time() - start
                if duration > 1.0:  # Log only slow steps (>1s)
                    # Use print for profiling info since this is outside class context
                    print(f"[PROFILE] {step_name} took {duration:.2f}s")
                return result
            except Exception as e:
                duration = time.time() - start
                # Use print for profiling error since this is outside class context
                print(f"[PROFILE] {step_name} failed after {duration:.2f}s: {e}")
                raise
        return wrapper
    return decorator



def _process_single_resume_worker(payload: Dict[str, Any]) -> ProcessingResult:
    """Worker function for processing a single resume in a separate process.
    
    Args:
        payload: Serializable dictionary containing file data and processing parameters
        
    Returns:
        ProcessingResult with processing outcome
    """
    start_time = time.time()
    filename = payload.get('filename', 'unknown')
    
    try:
        # Validate payload
        if 'file_content' not in payload:
            return ProcessingResult(
                success=False,
                filename=filename,
                error="Missing file content in payload",
                processing_time=time.time() - start_time
            )
        
        # Reconstruct file object from bytes
        file_obj = BytesIO(payload['file_content'])
        
        # Create processor and process resume
        processor = ResumeProcessor()
        file_data = {
            'filename': filename,
            'file': file_obj,
            'text': payload.get('text', ''),
            'recipient_email': payload.get('recipient_email', ''),
            'sender_email': payload.get('sender_email', ''),
            'sender_password': payload.get('sender_password', ''),
            'smtp_server': payload.get('smtp_server', ''),
            'smtp_port': payload.get('smtp_port', 465),
            'email_subject': payload.get('email_subject', ''),
            'email_body': payload.get('email_body', ''),
        }
        
        result = processor.process_single_resume(file_data)
        
        # Convert to ProcessingResult if it's not already
        if isinstance(result, dict):
            return ProcessingResult(
                success=result.get('success', False),
                filename=filename,
                error=result.get('error'),
                points_added=result.get('points_added', 0),
                processing_time=time.time() - start_time,
                tech_stacks_used=result.get('tech_stacks_used', []),
                modified_content=result.get('modified_content'),
                metadata=result.get('metadata', {})
            )
        return result
        
    except Exception as e:
        # Use print for worker error since this is outside class context
        print(f"Worker error processing {filename}: {e}")
        return ProcessingResult(
            success=False,
            filename=filename,
            error=str(e),
            processing_time=time.time() - start_time
        )


class ResumeProcessor:
    @with_file_error_handling("resume_processing")
    def process_single_resume(self, file_data: Dict[str, Any], progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """
        Process a single resume file with comprehensive error handling.
        
        Args:
            file_data: Dictionary containing file information and processing parameters
            progress_callback: Optional callback function for progress updates
            
        Returns:
            Dictionary with processing results
        """
        start_time = time.time()
        filename = file_data.get('filename', 'unknown')
        file = file_data.get('file')
        text = file_data.get('text', '')
        
        # Convert file to BytesIO
        if hasattr(file, 'read'):
            file_content = file.read()
            file.seek(0)  # Reset file pointer
            file_obj = BytesIO(file_content)
        else:
            file_obj = None
        
        # Create error context for detailed error handling
        error_context = ErrorContext(
            operation="resume_processing",
            severity=ErrorSeverity.HIGH,
            component="resume_processor",
            details={
                "resume_file": filename,  # Using a different name to avoid conflict
                "text_length": len(text) if text else 0,
                "has_file_obj": file_obj is not None
            }
        )
        
        try:
            if not file_obj:
                raise FileProcessingError(
                    "Missing file object in file_data",
                    filename
                )
                
            # Update progress if callback provided
            if progress_callback:
                progress_callback(f"Processing {filename}...")
                
            # Process the document
            from resume_customizer.processors.document_processor import get_document_processor
            doc_processor = get_document_processor()
            
            try:
                # Parse input text to get points
                from resume_customizer.parsers.text_parser import parse_input_text
                parsed_points, tech_stacks_used = parse_input_text(text)
                # Create the tuple that document processor expects
                parsed_data = (parsed_points, tech_stacks_used)
            except Exception as parser_error:
                # Create new error context with updated severity
                parser_error_context = ErrorContext(
                    operation=error_context.operation,
                    severity=ErrorSeverity.MEDIUM,
                    component=error_context.component,
                    details={**error_context.details, "error_stage": "parsing"}
                )
                raise FileProcessingError(
                    f"Failed to parse input text: {str(parser_error)}",
                    context=parser_error_context,
                    error_code="PARSER_ERROR",
                    original_exception=parser_error
                )

            # Extract matched companies if provided in file_data
            matched_companies = file_data.get('matched_companies') if isinstance(file_data, dict) else None

            try:
                # Process the document with the parsed points tuple
                processed_doc = doc_processor.process_document(file_obj, parsed_data, matched_companies=matched_companies)
            except Exception as doc_error:
                raise FileProcessingError(
                    f"Document processing failed: {str(doc_error)}",
                    context=error_context,
                    error_code="DOC_PROCESSING_ERROR",
                    original_exception=doc_error
                )
            
            # Prepare output bytes from processed document (handle BytesIO, file-like, Document, bytes)
            if isinstance(processed_doc, BytesIO):
                try:
                    processed_doc.seek(0)
                except Exception:
                    pass
                output_bytes = processed_doc.read()
            elif hasattr(processed_doc, 'read') and callable(getattr(processed_doc, 'read')):
                try:
                    if hasattr(processed_doc, 'seek'):
                        processed_doc.seek(0)
                except Exception:
                    pass
                data = processed_doc.read()
                output_bytes = data if isinstance(data, (bytes, bytearray)) else bytes(data)
            elif hasattr(processed_doc, "save"):
                tmp = BytesIO()
                processed_doc.save(tmp)
                tmp.seek(0)
                output_bytes = tmp.getvalue()
            elif isinstance(processed_doc, (bytes, bytearray)):
                output_bytes = bytes(processed_doc)
            else:
                raise TypeError(f"Unsupported processed document type: {type(processed_doc).__name__}")

            # Return success result
            return {
                'success': True,
                'points_added': len(parsed_points),
                'modified_content': output_bytes,
                'filename': filename,
                'processing_time': time.time() - start_time,
                'tech_stacks_used': tech_stacks_used,
                'metadata': {
                    'parser_used': 'TextParser',
                    'processing_diagnostics': FileErrorHandler.get_diagnostics(filename)
                }
            }
            
        except FileProcessingError as fpe:
            # This will be handled by the with_file_error_handling decorator
            # which provides user-friendly messages and detailed logging
            raise
            
        except Exception as e:
            # Log the error with enhanced context (without depending on FileErrorHandler static methods)
            error_details = {
                'filename': filename,
                'error_type': type(e).__name__,
                'error_message': str(e),
                'traceback': traceback.format_exc(),
                'timestamp': datetime.now().isoformat()
            }
            logging.error(f"Error processing resume {filename}: {error_details}")
            
            # Prepare diagnostics with graceful fallback
            try:
                diagnostics = FileErrorHandler.get_diagnostics(filename)
            except Exception:
                diagnostics = {'filename': filename}
            
            # Return error result with enhanced diagnostics
            return {
                'success': False,
                'error': str(e),
                'error_details': error_details,
                'filename': filename,
                'processing_time': time.time() - start_time,
                'diagnostics': diagnostics
            }
    
    def process_single_resume_async(self, file_data: Dict[str, Any]):
        """
        Submit a resume processing job to Celery and return the AsyncResult.
        Serializes file content to base64 for JSON compatibility.
        """
        try:
            # Use Celery app to get the registered task instead of direct import
            # This avoids circular import issues
            from infrastructure.async_processing.tasks import celery_app
            process_resume_task = celery_app.tasks.get('tasks.process_resume_task')
            
            if process_resume_task is None:
                # Try to import the tasks module to register tasks
                try:
                    from infrastructure.async_processing import tasks  # This should register the task
                    process_resume_task = celery_app.tasks.get('tasks.process_resume_task')
                except ImportError:
                    pass
            
            if process_resume_task is None:
                raise ImportError("Task 'tasks.process_resume_task' not found in Celery app")
                
        except ImportError as import_error:
            import os
            import sys
            cwd = os.getcwd()
            python_path = sys.path[:3]  # Show first 3 paths
            error_msg = f"""Celery tasks module not found. Debugging info:
            - Current working directory: {cwd}
            - Python path (first 3): {python_path}
            - Import error: {str(import_error)}
            - Make sure tasks.py exists and Celery worker is running
            - Try running from project root directory"""
            raise RuntimeError(error_msg)
        
        try:
            # Create a copy of file_data to avoid modifying the original
            serialized_data = file_data.copy()
            
            # Serialize the file object to base64 for JSON compatibility
            if 'file' in serialized_data and serialized_data['file'] is not None:
                file_obj = serialized_data['file']
                
                # Read file content
                file_obj.seek(0)  # Ensure we're at the start of the file
                file_content = file_obj.read()
                file_obj.seek(0)  # Reset position for potential future use
                
                # Convert to base64 string
                import base64
                file_content_b64 = base64.b64encode(file_content).decode('utf-8')
                
                # Replace file object with base64 content
                serialized_data['file_content_b64'] = file_content_b64
                del serialized_data['file']  # Remove non-serializable file object
            else:
                raise ValueError("File object is required for async processing")
            
            return process_resume_task.delay(serialized_data)
            
        except Exception as e:
            error_msg = str(e)
            if 'Redis' in error_msg or 'Connection refused' in error_msg:
                raise RuntimeError("Celery broker (Redis) connection failed. Make sure Redis is running or configure an alternative broker.")
            elif 'NoneType' in error_msg and 'Redis' in error_msg:
                raise RuntimeError("Redis configuration issue. Check celeryconfig.py and ensure Redis is properly installed.")
            elif 'not JSON serializable' in error_msg:
                raise RuntimeError("File serialization failed. Please try again or use synchronous processing.")
            else:
                raise RuntimeError(f"Celery task submission failed: {error_msg}")

    def get_async_result(self, task_id: str):
        """
        Get the result/status of a Celery async task by task_id.
        """
        try:
            from infrastructure.async_processing.tasks import celery_app
        except ImportError:
            raise RuntimeError("Celery config not found. Make sure celeryconfig.py exists and Celery is installed.")
        
        try:
            async_result = celery_app.AsyncResult(task_id)
            
            # Get the state safely
            state = async_result.state
            
            # Get result safely, handling different states
            result = None
            if async_result.ready():
                try:
                    result = async_result.result
                except Exception as result_error:
                    # If result retrieval fails, include error info
                    result = {
                        'error': f"Result retrieval failed: {str(result_error)}",
                        'success': False
                    }
            
            # Get additional task info if available
            task_info = {
                'state': state,
                'result': result
            }
            
            # Add progress info if available (for PROGRESS state)
            if state == 'PROGRESS' and hasattr(async_result, 'info') and async_result.info:
                task_info['info'] = async_result.info
            
            return task_info
            
        except Exception as e:
            error_msg = str(e)
            
            # Handle specific error types
            if 'NotRegistered' in error_msg:
                return {
                    'state': 'FAILURE',
                    'result': {
                        'error': 'Task not found or not registered with worker. Make sure the Celery worker is running.',
                        'success': False
                    }
                }
            elif 'Redis' in error_msg or 'Connection refused' in error_msg:
                return {
                    'state': 'FAILURE', 
                    'result': {
                        'error': 'Redis connection failed. Check if Redis server is running.',
                        'success': False
                    }
                }
            else:
                return {
                    'state': 'FAILURE',
                    'result': {
                        'error': f"Failed to get task status: {error_msg}",
                        'success': False
                    }
                }
    @rate_limit("batch_processing", limit=5, window=600)  # 5 batch operations per 10 minutes
    def process_all_resumes(self, files: List[Any]) -> List[Dict[str, Any]]:
        """
        Process all resumes for bulk preview/generation (no email).
        Returns a list of preview/result dicts for each file.
        """
        results = []
        for file in files:
            # Try to get file name and text from session state if available
            file_name = getattr(file, 'name', 'Resume')
            resume_inputs = st.session_state.get('resume_inputs', {})
            file_data = resume_inputs.get(file_name, {})
            text = file_data.get('text', '')
            manual_text = file_data.get('manual_text', '')
            # Use PreviewGenerator for consistent preview output
            try:
                # PreviewGenerator is defined in this same file
                previewer = PreviewGenerator()
                preview_result = previewer.generate_preview(file, text, manual_text)
                preview_result['file_name'] = file_name
                results.append(preview_result)
            except Exception as e:
                results.append({'file_name': file_name, 'success': False, 'error': str(e)})
        return results
    """Main processor for individual resume operations."""
    
class RobustResumeProcessor:
    """Enhanced resume processor with robust error handling and recovery."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
        self.error_recovery = get_error_recovery_manager()
        self._metrics = get_analytics_manager()
        self._retries = 0
        self._max_retries = 3
    
    @record_metrics("robust_processing")
    def process_with_recovery(self, content: BytesIO, tech_stack: str, 
                            error_context: Optional[ErrorContext] = None) -> BytesIO:
        """Process resume with error recovery mechanisms."""
        try:
            # Parse tech stack with fallback options
            try:
                parsed_points = parse_input_text_restricted(tech_stack)
            except RestrictedFormatError:
                # Fallback to legacy parser if restricted format fails
                parsed_points = parse_input_text(tech_stack)
            
            # Process document with recovery
            for attempt in range(self._max_retries):
                try:
                    processed_doc = self.doc_processor.process_document(
                        content,
                        parsed_points
                    )
                    return processed_doc
                except Exception as e:
                    if self.error_recovery and attempt < self._max_retries - 1:
                        self.error_recovery.attempt_recovery(error_context or ErrorContext(
                            severity=ErrorSeverity.MEDIUM,
                            component="robust_processor",
                            operation="process_document",
                            details={"attempt": attempt + 1}
                        ))
                        self._retries += 1
                        continue
                    raise
                    
        except Exception as e:
            # Use print for robust processing error since this is outside class context
            print(f"Robust processing failed: {str(e)}")
            if error_context:
                error_context.details["final_error"] = str(e)
            raise

    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics."""
        return {
            "retries": self._retries,
            "success_rate": self._metrics.get_summary("robust_processing.success")
                if self._metrics else None
        }

    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
        self.file_processor = FileProcessor()
        # Use enhanced robust processor for error recovery
        self.robust_processor = RobustResumeProcessor()
        # Disabled distributed cache - monitoring removed
        self.cache_manager = None
        # Progress tracking
        self.progress_tracker = get_progress_tracker()
        # Legacy support - Advanced cache: (filename, text, manual_text) -> result
        self._result_cache: Dict[str, Dict[str, Any]] = {}
        self._cache_lock = threading.Lock()
        self._audit_log = []
        self._task_queue = queue.Queue()
        self._worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self._worker_thread.start()

    def _worker_loop(self):
        while True:
            try:
                task, args, kwargs = self._task_queue.get()
                try:
                    task(*args, **kwargs)
                except Exception as e:
                    # Use print for background task error since this is outside class context
                    print(f"Background task error: {e}")
            except Exception:
                pass
        
    @record_metrics("resume_processing")
    # Removed circuit breaker and cache decorators - monitoring disabled
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Process a single resume with enhanced error recovery, caching, and monitoring.
        
        Args:
            file_data: Dictionary containing file information and settings
            progress_callback: Optional callback for progress updates
            
        Returns:
            Result dictionary with processing status and data
        """
        start_time = time.time()
        cache_key = f"{file_data['filename']}|{file_data.get('text','')}|{file_data.get('manual_text','')}"
        with self._cache_lock:
            if cache_key in self._result_cache:
                self.logger.info(f"[CACHE] Hit for {file_data['filename']}")
                return self._result_cache[cache_key]
        try:
            filename = file_data['filename']
            file_obj = file_data['file']
            raw_text = file_data['text']
            user = None
            try:
                user = st.session_state.get('user_id', 'unknown')
            except Exception:
                user = 'unknown'
            audit_logger.log(
                action="process_single_resume",
                user=user,
                details={"filename": filename},
                status="started"
            )
            if progress_callback:
                progress_callback(f"Parsing tech stacks for {filename}...")
            # Parse tech stacks and points (now backed by LRU cache in parser)
            manual_text = file_data.get('manual_text', '')
            
            # Use new restricted parser first, with fallback to legacy parser
            selected_points, tech_stacks_used = [], []
            parsing_method = "unknown"
            
            try:
                # Try restricted parser first
                selected_points, tech_stacks_used = parse_input_text_restricted(raw_text, manual_text)
                parsing_method = "restricted"
                
                if not selected_points or not tech_stacks_used:
                    # Empty results from restricted parser, try legacy fallback
                    self.logger.warning(f"Restricted parser returned empty results for {filename}, trying legacy parser")
                    from resume_customizer.parsers.text_parser import parse_input_text
                    selected_points, tech_stacks_used = parse_input_text(raw_text, manual_text)
                    parsing_method = "legacy_fallback"
                    
            except RestrictedFormatError as e:
                # Format error from restricted parser, try legacy fallback
                self.logger.warning(f"Restricted parser failed for {filename}: {str(e)}, trying legacy parser")
                try:
                    from resume_customizer.parsers.text_parser import parse_input_text
                    selected_points, tech_stacks_used = parse_input_text(raw_text, manual_text)
                    parsing_method = "legacy_fallback"
                except Exception as fallback_error:
                    self.logger.error(f"Both parsers failed for {filename}: {str(fallback_error)}")
                    result = {
                        'success': False,
                        'error': ERROR_MESSAGES.get("parsing_exception", "Parsing error: {error}").format(error=str(e)),
                        'filename': filename
                    }
                    audit_logger.log(
                        action="process_single_resume",
                        user=user,
                        details={"filename": filename},
                        status="failed",
                        error=result['error']
                    )
                    with self._cache_lock:
                        self._result_cache[cache_key] = result
                    return result
                    
            # Final check for empty results after all parsing attempts
            if not selected_points or not tech_stacks_used:
                result = {
                    'success': False,
                    'error': ERROR_MESSAGES.get("empty_parse_results", "No valid points detected in your input"),
                    'filename': filename
                }
                audit_logger.log(
                    action="process_single_resume",
                    user=user,
                    details={"filename": filename},
                    status="failed",
                    error=result['error']
                )
                with self._cache_lock:
                    self._result_cache[cache_key] = result
                return result
                
            self.logger.info(f"Successfully parsed {len(selected_points)} points using {parsing_method} method for {filename}")
            if progress_callback:
                progress_callback(f"Processing document for {filename}...")
            # Load and process document
            doc = Document(file_obj)
            
            # Detect document-wide bullet marker for consistency
            document_marker = self.doc_processor.bullet_formatter.detect_document_bullet_marker(doc)
            self.logger.info(f"Document bullet marker detected: '{document_marker}'")
            
            projects_data = self.doc_processor.project_detector.find_projects(doc)
            if not projects_data:
                result = {
                    'success': False,
                    'error': ERROR_MESSAGES["no_projects"].format(filename=filename),
                    'filename': filename
                }
                audit_logger.log(
                    action="process_single_resume",
                    user=user,
                    details={"filename": filename},
                    status="failed",
                    error=result['error']
                )
                with self._cache_lock:
                    self._result_cache[cache_key] = result
                return result
            
            # Convert ProjectInfo objects to structured format
            projects = []
            for i, project_info in enumerate(projects_data):
                projects.append({
                    'title': project_info.name,
                    'index': i,
                    'responsibilities_start': project_info.start_index,
                    'responsibilities_end': project_info.end_index,
                    'project_info': project_info  # Keep reference to original object
                })
            
            # Distribute and add points using round-robin logic
            distribution_result = self.doc_processor.point_distributor.distribute_points(
                projects, (selected_points, tech_stacks_used)
            )
            if not distribution_result.distribution:
                error_result = {
                    'success': False,
                    'error': ERROR_MESSAGES.get("point_distribution_failed", "Cannot distribute points because no valid tech stack data was found"),
                    'filename': filename,
                    'debug_info': {
                        'selected_points_count': len(selected_points),
                        'tech_stacks_count': len(tech_stacks_used),
                        'projects_count': len(projects),
                        'parsing_method': parsing_method
                    }
                }
                self.logger.error(f"Point distribution failed for {filename}. Points: {len(selected_points)}, Stacks: {len(tech_stacks_used)}, Projects: {len(projects)}")
                with self._cache_lock:
                    self._result_cache[cache_key] = error_result
                return error_result
            
            # Add points to each project with consistent formatting
            total_added = 0
            
            for project_name, points in distribution_result.distribution.items():
                # Find the corresponding project
                project = next((p['project_info'] for p in projects if p['title'] == project_name), None)
                if project and points:
                    added = self.doc_processor._add_points_to_project(doc, project, points, document_marker)
                    total_added += added
                    self.logger.debug(f"Added {added} points to project '{project_name}' with marker '{document_marker}'")
            
            points_added = total_added
            if progress_callback:
                progress_callback(f"Saving document for {filename}...")
            # Save to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            result = {
                'success': True,
                'filename': filename,
                'buffer': output_buffer.getvalue(),
                'tech_stacks': tech_stacks_used,
                'points_added': points_added,
                'email_data': self._extract_email_data(file_data)
            }
            audit_logger.log(
                action="process_single_resume",
                user=user,
                details={"filename": filename},
                status="success"
            )
            with self._cache_lock:
                self._result_cache[cache_key] = result
            return result
        except Exception as e:
            # Retry logic: up to 2 more times for transient errors
            for attempt in range(2):
                try:
                    time.sleep(0.5 * (attempt + 1))
                    return self.process_single_resume(file_data, progress_callback)
                except Exception:
                    continue
            self.logger.error(f"Error processing resume {file_data.get('filename','?')}: {e}")
            result = {'success': False, 'filename': file_data.get('filename', '?'), 'error': str(e)}
            audit_logger.log(
                action="process_single_resume",
                user=user if 'user' in locals() else 'unknown',
                details={"filename": file_data.get('filename', '?')},
                status="failed",
                error=str(e)
            )
            with self._cache_lock:
                self._result_cache[cache_key] = result
            return result
    
    def _extract_email_data(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract email configuration from file data.
        
        Args:
            file_data: File data dictionary
            
        Returns:
            Email configuration dictionary
        """
        return {
            'recipient': file_data.get('recipient_email', ''),
            'sender': file_data.get('sender_email', ''),
            'password': file_data.get('sender_password', ''),
            'smtp_server': file_data.get('smtp_server', ''),
            'smtp_port': file_data.get('smtp_port', 465),
            'subject': file_data.get('email_subject', ''),
            'body': file_data.get('email_body', '')
        }


class BulkResumeProcessor:
    """Handles bulk processing of multiple resumes with parallel execution."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.file_processor = FileProcessor()
        self.email_manager = get_email_manager()
        # Determine a sensible default for max workers based on CPU
        try:
            import os
            cpu_count = os.cpu_count() or 4
            # Use min to avoid oversubscription; ThreadPool benefits from I/O overlap but docx is CPU-heavy
            self._default_workers = max(2, min(8, cpu_count))
        except Exception:
            self._default_workers = 4
    
    def process_resumes_bulk(
        self,
        files_data: List[Dict[str, Any]],
        max_workers: int = None,
        progress_callback: Optional[Callable] = None,
        use_process_pool: Optional[bool] = None,
        memory_limit_mb: int = 1024,
        ui_update_interval: float = 0.5
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Process multiple resumes in parallel.
        
        Args:
            files_data: List of file data dictionaries
            max_workers: Maximum number of parallel workers
            progress_callback: Optional progress callback
            
        Returns:
            Tuple of (processed_resumes, failed_resumes)
        """
        import time, psutil, os
        processed_resumes = []
        failed_resumes = []
        last_ui_update = time.time()
        if max_workers is None:
            cpu_count = os.cpu_count() or 4
            # For CPU-bound, use cpu_count; for I/O-bound, allow more
            max_workers = min(max(cpu_count, 4), len(files_data), 16)
        if max_workers < 1:
            max_workers = 1

        if use_process_pool is None:
            use_process_pool = len(files_data) >= 3 and (os.cpu_count() or 2) > 2

        # Ensure all files have proper names before processing
        optimized_files_data = []
        for file_data in files_data:
            file_data['file'] = self.file_processor.ensure_file_has_name(
                file_data['file'], file_data['filename']
            )
            optimized_files_data.append(file_data)

        def memory_check():
            try:
                process = psutil.Process(os.getpid())
                mem_mb = process.memory_info().rss / 1024 / 1024
                return mem_mb > memory_limit_mb
            except Exception:
                return False

        def throttled_progress_callback(msg):
            nonlocal last_ui_update
            now = time.time()
            if progress_callback and (now - last_ui_update > ui_update_interval):
                progress_callback(msg)
                last_ui_update = now

        if use_process_pool:
            # Only send minimal data: file path, not bytes
            # If file objects are not file paths, fallback to bytes
            payloads = []
            for file_data in optimized_files_data:
                file_obj = file_data['file']
                if hasattr(file_obj, 'name') and os.path.exists(file_obj.name):
                    payloads.append({
                        'filename': file_data['filename'],
                        'file_path': file_obj.name,
                        'text': file_data['text'],
                        'recipient_email': file_data.get('recipient_email', ''),
                        'sender_email': file_data.get('sender_email', ''),
                        'sender_password': file_data.get('sender_password', ''),
                        'smtp_server': file_data.get('smtp_server', ''),
                        'smtp_port': file_data.get('smtp_port', 465),
                        'email_subject': file_data.get('email_subject', ''),
                        'email_body': file_data.get('email_body', ''),
                    })
                else:
                    # fallback to bytes
                    file_bytes = file_obj.getvalue() if hasattr(file_obj, 'getvalue') else file_obj.read()
                    payloads.append({
                        'filename': file_data['filename'],
                        'file_content': file_bytes,
                        'text': file_data['text'],
                        'recipient_email': file_data.get('recipient_email', ''),
                        'sender_email': file_data.get('sender_email', ''),
                        'sender_password': file_data.get('sender_password', ''),
                        'smtp_server': file_data.get('smtp_server', ''),
                        'smtp_port': file_data.get('smtp_port', 465),
                        'email_subject': file_data.get('email_subject', ''),
                        'email_body': file_data.get('email_body', ''),
                    })

            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {executor.submit(_process_single_resume_worker, p): p['filename'] for p in payloads}
                for future in concurrent.futures.as_completed(future_to_file):
                    filename = future_to_file[future]
                    try:
                        result = future.result()
                        if result['success']:
                            processed_resumes.append(result)
                        else:
                            failed_resumes.append(result)
                        throttled_progress_callback(f"Processed {filename}")
                        if memory_check():
                            break
                    except Exception as e:
                        failed_resumes.append({'success': False, 'filename': filename, 'error': str(e)})
        else:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {
                    executor.submit(self.resume_processor.process_single_resume, file_data, throttled_progress_callback):
                    file_data['filename'] for file_data in optimized_files_data
                }
                for future in concurrent.futures.as_completed(future_to_file):
                    filename = future_to_file[future]
                    try:
                        result = future.result()
                        if result['success']:
                            processed_resumes.append(result)
                        else:
                            failed_resumes.append(result)
                        throttled_progress_callback(f"Processed {filename}")
                        if memory_check():
                            break
                    except Exception as e:
                        failed_resumes.append({'success': False, 'filename': filename, 'error': str(e)})

        self.file_processor.cleanup_memory()
        return processed_resumes, failed_resumes
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """
        Send emails for processed resumes.
        
        Args:
            processed_resumes: List of processed resume data
            progress_callback: Optional progress callback
            
        Returns:
            List of email sending results
        """
        return self.email_manager.send_batch_emails(processed_resumes, progress_callback)


class PreviewGenerator:
    """Generates preview of resume changes without modifying the original."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
    
    def generate_preview(
        self,
        file_obj,
        input_text: str,
        manual_text: str = "",
        matched_companies: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Generate a preview of changes that will be made to the resume.
        
        Args:
            file_obj: Resume file object
            input_text: Tech stack input text
            manual_text: Optional manual points text
            
        Returns:
            Preview result dictionary
        """
        try:
            # Use the same parser as the main processing for consistency
            selected_points, tech_stacks_used = parse_input_text(input_text, manual_text)

            if not selected_points or not tech_stacks_used:
                return {
                    'success': False,
                    'error': 'Could not parse tech stacks from input. Please check the format.'
                }

            # Load document and find projects
            doc = Document(file_obj)

            # Detect document-wide bullet marker for consistency
            document_marker = self.doc_processor.bullet_formatter.detect_document_bullet_marker(doc)

            projects_data = self.doc_processor.project_detector.find_projects(doc)

            if not projects_data:
                return {
                    'success': False,
                    'error': 'No projects with Responsibilities sections found'
                }

            # Convert ProjectInfo objects to structured format
            projects = []
            for i, project_info in enumerate(projects_data):
                projects.append({
                    'title': project_info.name,
                    'index': i,
                    'responsibilities_start': project_info.start_index,
                    'responsibilities_end': project_info.end_index,
                    'project_info': project_info
                })

            # Create preview document copy
            temp_buffer = BytesIO()
            doc.save(temp_buffer)
            temp_buffer.seek(0)

            preview_doc = Document(temp_buffer)
            preview_projects_data = self.doc_processor.project_detector.find_projects(preview_doc)

            # Convert to structured format
            preview_projects = []
            for i, project_info in enumerate(preview_projects_data):
                preview_projects.append({
                    'title': project_info.name,
                    'index': i,
                    'responsibilities_start': project_info.start_index,
                    'responsibilities_end': project_info.end_index,
                    'project_info': project_info
                })

            # Apply changes to preview - pass the raw ProjectInfo objects like in document processing
            try:
                from config import PARSING_CONFIG
                threshold = PARSING_CONFIG.get('company_match_threshold', 70)
            except Exception:
                threshold = 70

            distribution_result = self.doc_processor.point_distributor.distribute_points(
                preview_projects_data, (selected_points, tech_stacks_used), company_priority=matched_companies, company_match_threshold=threshold
            )
            if not distribution_result.distribution:
                return {
                    'success': False, 
                    'error': ERROR_MESSAGES.get("point_distribution_failed", "Cannot distribute points because no valid tech stack data was found"),
                    'debug_info': {
                        'selected_points_count': len(selected_points),
                        'tech_stacks_count': len(tech_stacks_used),
                        'projects_count': len(preview_projects)
                    }
                }
            
            # Add points to each project with consistent formatting
            total_added = 0
            project_points_mapping = {}
            
            for project_name, points in distribution_result.distribution.items():
                # Find the corresponding project by name
                project = next((p for p in preview_projects_data if p.name == project_name), None)
                if project and points:
                    # Store points mapping for display
                    project_points_mapping[project_name] = {
                        'points': [p.get('text', str(p)) for p in points]
                    }
                    
                    # Add points with consistent formatting
                    added = self.doc_processor._add_points_to_project(preview_doc, project, points, document_marker)
                    total_added += added
            
            points_added = total_added
            
            # Generate preview content
            preview_content = self._extract_preview_content(preview_doc)
            
            # Extract original content for comparison
            original_content = self._extract_preview_content(doc)
            
            return {
                'success': True,
                'points_added': points_added,
                'tech_stacks_used': tech_stacks_used,
                'selected_points': selected_points,
                'preview_content': preview_content,
                'original_content': original_content,
                'preview_doc': preview_doc,
                'projects_count': len(projects),
                'project_points_mapping': project_points_mapping,
                'document_marker': document_marker,
                'last_updated': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_preview_content(self, doc: Document) -> str:
        """
        Extract readable content from document for preview with enhanced formatting.
        
        Args:
            doc: Document to extract content from
            
        Returns:
            Text content of the document with enhanced formatting
        """
        content_parts = []
        in_project = False
        current_project = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Detect project headers (typically bold text)
            is_header = any(run.bold for run in para.runs if run.text.strip())
            
            if is_header and len(text) < 100:  # Likely a header/title
                in_project = True
                current_project = text
                content_parts.append(f"\n== {text} ==")
            elif in_project and text.startswith(('•', '-', '●', '○', '■', '▪', '▫', '⦿')):
                # Format bullet points with indentation for better readability
                content_parts.append(f"  {text}")
            else:
                content_parts.append(text)
        
        return "\n".join(content_parts)


class ResumeManager:
    def process_all_resumes(self, files: List[Any]) -> List[Dict[str, Any]]:
        """Process all resumes for bulk preview/generation (no email)."""
        return self.resume_processor.process_all_resumes(files)
    """Main manager class that coordinates all resume processing operations."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.bulk_processor = BulkResumeProcessor()
        self.preview_generator = PreviewGenerator()
        self.email_manager = get_email_manager()
    
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Process a single resume."""
        return self.resume_processor.process_single_resume(file_data, progress_callback)
    
    def process_bulk_resumes(
        self, 
        files_data: List[Dict[str, Any]], 
        max_workers: int = None,
        progress_callback: Optional[Callable] = None
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process multiple resumes in parallel."""
        return self.bulk_processor.process_resumes_bulk(files_data, max_workers, progress_callback)
    
    # Async integration (Celery) delegations
    def process_single_resume_async(self, file_data: Dict[str, Any]):
        """Submit a resume processing job to Celery and return the AsyncResult."""
        return self.resume_processor.process_single_resume_async(file_data)

    def get_async_result(self, task_id: str) -> Dict[str, Any]:
        """Fetch status/result for an async Celery task by ID."""
        return self.resume_processor.get_async_result(task_id)

    def generate_preview(
        self, 
        file_obj, 
        input_text: str, 
        manual_text: str = "",
        matched_companies: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Generate preview of resume changes."""
        return self.preview_generator.generate_preview(file_obj, input_text, manual_text, matched_companies=matched_companies)
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """Send batch emails for processed resumes."""
        return self.bulk_processor.send_batch_emails(processed_resumes, progress_callback)
    
    def send_single_email(
        self, 
        smtp_server: str,
        smtp_port: int,
        sender_email: str,
        sender_password: str,
        recipient_email: str,
        subject: str,
        body: str,
        attachment_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """Send a single email."""
        return self.email_manager.send_single_email(
            smtp_server, smtp_port, sender_email, sender_password,
            recipient_email, subject, body, attachment_data, filename
        )
    
    def validate_email_config(self, email_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate email configuration."""
        return self.email_manager.validate_email_config(email_data)
    
    def cleanup(self) -> None:
        """Clean up resources and connections."""
        self.email_manager.close_all_connections()


# Global resume manager accessor
# Use a lightweight proxy so callers importing `resume_manager` don't hold onto a stale instance
class _ResumeManagerProxy:
    def __getattr__(self, name):
        return getattr(get_resume_manager(), name)

resume_manager = _ResumeManagerProxy()


@st.cache_resource
def get_resume_manager(_version: str = "v2.3") -> ResumeManager:
    """
    Get the global resume manager instance.
    
    Returns:
        ResumeManager instance
    """
    # Force cache refresh by creating new instance
    return ResumeManager()


