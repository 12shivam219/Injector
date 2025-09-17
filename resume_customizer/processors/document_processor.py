"""
Document processing module for Resume Customizer application.
Handles Word document operations, project detection, and point insertion.
"""

import time
from typing import List, Dict, Any, Tuple, Optional
from io import BytesIO
from docx import Document

import logging
from infrastructure.utilities.logger import get_logger
from infrastructure.utilities.structured_logger import get_structured_logger
from ..detectors.project_detector import ProjectDetector, ProjectInfo
from ..formatters.bullet_formatter import BulletFormatter, BulletFormatting
from .point_distributor import PointDistributor

# Initialize logger with fallback
try:
    logger = get_logger()
except Exception:
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)

try:
    structured_logger = get_structured_logger("document_processor")
except Exception:
    structured_logger = logging.getLogger("document_processor")
    structured_logger.setLevel(logging.INFO)


class DocumentProcessor:
    """Handles document processing operations with enhanced formatting preservation."""
    
    def __init__(self):
        self.project_detector = ProjectDetector()
        self.bullet_formatter = BulletFormatter()
        self.point_distributor = PointDistributor()
    
    def process_document(self, file_content: BytesIO, parsed_points: Tuple[List[str], List[str]], matched_companies: Optional[List[str]] = None) -> BytesIO:
        """
        Process a document by adding tech stack points to projects.
        
        Args:
            file_content: Document content as BytesIO
            parsed_points: Tuple of (selected_points, tech_stacks_used)
            
        Returns:
            Processed document as BytesIO
        """
        try:
            # Load document
            doc = Document(file_content)
            
            # Detect document-wide bullet marker for consistency
            document_marker = self.bullet_formatter.detect_document_bullet_marker(doc)
            logger.info(f"Detected document bullet marker: '{document_marker}'")
            
            # Find projects in document
            projects = self.project_detector.find_projects(doc)
            if not projects:
                raise ValueError("No projects found in document")
            
            # Distribute points to projects (prefer matched companies if provided)
            # Read threshold from global parsing config if available
            try:
                from config import PARSING_CONFIG
                threshold = PARSING_CONFIG.get('company_match_threshold', 70)
            except Exception:
                threshold = 70

            distribution_result = self.point_distributor.distribute_points(projects, parsed_points, company_priority=matched_companies, company_match_threshold=threshold)
            
            # Add points to projects
            total_added = 0
            for project_name, points in distribution_result.distribution.items():
                project = next((p for p in projects if p.name == project_name), None)
                if project and points:
                    added = self._add_points_to_project(doc, project, points, document_marker)
                    total_added += added
            
            # Save to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            logger.info(f"Document processed successfully, added {total_added} points")
            return output_buffer
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise
    
    def _add_points_to_project(self, doc: Document, project: ProjectInfo, 
                               points: List[Dict[str, Any]], document_marker: str) -> int:
        """
        Add points to a specific project after existing bullet points.
        """
        if not points:
            return 0

        try:
            # Find the last bullet paragraph in the project
            last_bullet_index = None
            for i in range(project.start_index, project.end_index + 1):
                para = doc.paragraphs[i]
                if self.bullet_formatter._is_bullet_point(para.text):
                    last_bullet_index = i
            
            # If no bullets exist, find where to insert them
            if last_bullet_index is None:
                # Look for "Responsibilities" or similar heading
                for i in range(project.start_index, project.end_index + 1):
                    para = doc.paragraphs[i]
                    if any(keyword in para.text.lower() for keyword in ["responsibilities", "duties", "achievements"]):
                        last_bullet_index = i
                        break
                
                # If still not found, find the first bullet-like paragraph after the role line
                if last_bullet_index is None:
                    # First try to find any existing bullet points in the project
                    for i in range(project.start_index, project.end_index + 1):
                        para = doc.paragraphs[i]
                        text = para.text.strip()
                        # Check for common bullet markers or dash at the beginning
                        if text.startswith('•') or text.startswith('-') or text.startswith('*') or text.startswith('○'):
                            last_bullet_index = i
                            break
                
                # If still not found and project has a role, find the role line and add after it
                if last_bullet_index is None and project.role:
                    role_index = None
                    # Find the role line
                    for i in range(project.start_index, project.end_index + 1):
                        if project.role in doc.paragraphs[i].text:
                            role_index = i
                            break
                    
                    # If role found, look for the first non-empty paragraph after it
                    if role_index is not None:
                        for i in range(role_index + 1, project.end_index + 1):
                            if doc.paragraphs[i].text.strip():
                                last_bullet_index = i - 1  # Insert before this non-empty paragraph
                                break
                        
                        # If no non-empty paragraph found, use the role line
                        if last_bullet_index is None:
                            last_bullet_index = role_index
                
                # If still not found, use the line after the company/date header
                if last_bullet_index is None:
                    last_bullet_index = project.start_index
                    # Ensure we're not out of bounds
                    if last_bullet_index >= len(doc.paragraphs):
                        last_bullet_index = project.end_index
            
            insertion_para = doc.paragraphs[last_bullet_index]

            # Get bullet formatting (fallback to document marker)
            existing_formatting = self._get_project_bullet_formatting(doc, project, document_marker)
            fallback_formatting = BulletFormatting(
                runs_formatting=[],
                paragraph_formatting={},
                style="Normal",
                bullet_marker=existing_formatting.bullet_marker if existing_formatting else document_marker,
                bullet_separator=" ",
                list_format={"is_list": False}
            )

            points_added = 0
            for point_data in points:
                point_text = point_data.get('text', str(point_data))

                # Create new paragraph and insert after the first bullet
                new_para = doc.add_paragraph()
                insertion_para._element.addnext(new_para._element)

                # Update insertion_para to the newly inserted paragraph (next point goes after it)
                insertion_para = new_para

                # Apply formatting
                self.bullet_formatter.apply_formatting(
                    insertion_para,
                    existing_formatting,
                    point_text,
                    fallback_formatting=fallback_formatting
                )

                points_added += 1

            # Update project's end_index to include new points
            project.end_index += points_added

            logger.info(
                f"Added {points_added} points to project '{project.name}' "
                f"after first bullet with marker '{existing_formatting.bullet_marker}'"
            )
            return points_added

        except Exception as e:
            logger.error(f"Failed to add points to project '{project.name}': {e}")
            return 0
    
    def _get_project_bullet_formatting(self, doc: Document, project: ProjectInfo, 
                                       document_marker: str) -> BulletFormatting:
        """
        Get bullet formatting from existing project bullets or fallback to document-wide marker.
        """
        for i in range(project.start_index, min(project.end_index + 1, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if self.bullet_formatter._is_bullet_point(para.text):
                formatting = self.bullet_formatter.extract_formatting(doc, i)
                if formatting:
                    if not formatting.bullet_marker:
                        formatting.bullet_marker = document_marker
                    logger.debug(
                        f"Extracted formatting from project '{project.name}' using marker '{formatting.bullet_marker}'"
                    )
                    return formatting

        # Fallback to document marker
        return BulletFormatting(
            runs_formatting=[{
                'text': '',
                'font_name': None,
                'font_size': None,
                'bold': None,
                'italic': None,
                'underline': None,
                'color': None
            }],
            paragraph_formatting={},
            style='Normal',
            bullet_marker=document_marker,
            bullet_separator=' ',
            list_format={
                'ilvl': 0,
                'numId': 1,
                'style': 'List Bullet',
                'indent': 0,
                'is_list': True
            }
        )


class FileProcessor:
    """Handles file operations and memory management."""
    
    def __init__(self):
        self.processed_files = {}
    
    def ensure_file_has_name(self, file_obj, filename: str):
        """Ensure file object has a name attribute."""
        if not hasattr(file_obj, 'name'):
            file_obj.name = filename
        return file_obj
    
    def cleanup_memory(self):
        """Clean up processed file cache."""
        self.processed_files.clear()


# Global document processor instance
_document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get singleton document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
