"""
Enhanced Resume Processor with Format Matching
"""

from typing import Dict, Any, Optional, List
import docx
from datetime import datetime

from database.models import ResumeFormat
from resume_customizer.analyzers.format_analyzer import FormatAnalyzer
from infrastructure.utilities.logger import get_logger

logger = get_logger()

class FormatAwareProcessor:
    """
    Enhances resume processing with format pattern matching
    """
    
    def __init__(self, db_session):
        self.session = db_session
        self.analyzer = FormatAnalyzer()
        
    def process_resume(self, doc: docx.Document) -> Dict[str, Any]:
        """
        Process a resume using format matching for better accuracy
        
        Args:
            doc: The Word document to process
            
        Returns:
            Processing results including matched format
        """
        try:
            # Find best matching format
            best_format = None
            best_score = 0
            matched_elements = {}
            
            formats = self.session.query(ResumeFormat).all()
            
            for fmt in formats:
                score, elements = self.analyzer.match_format(doc, fmt)
                if score > best_score:
                    best_score = score
                    best_format = fmt
                    matched_elements = elements
            
            if best_format and best_score >= 40:  # 40% minimum match threshold
                # Record the match
                doc_hash = self.analyzer._generate_hash(
                    "\n".join(p.text for p in doc.paragraphs)
                )
                
                match = ResumeFormatMatch(
                    format_id=best_format.id,
                    resume_hash=doc_hash,
                    match_score=best_score,
                    matched_elements=matched_elements
                )
                
                self.session.add(match)
                best_format.match_count += 1
                best_format.last_used = datetime.now()
                self.session.commit()
                
                # Use the matched format for processing
                return self._process_with_format(doc, best_format, matched_elements)
            else:
                # Fallback to default processing if no good match
                return self._process_without_format(doc)
                
        except Exception as e:
            logger.error(f"Error in format-aware processing: {str(e)}")
            raise

    def _process_with_format(self, 
                           doc: docx.Document, 
                           format: ResumeFormat, 
                           matched: Dict) -> Dict[str, Any]:
        """
        Process a resume using a matched format pattern
        """
        try:
            # Extract sections using format patterns
            sections = self._extract_sections_with_format(doc, format)
            
            # Extract companies using format patterns
            companies = self._extract_companies_with_format(doc, format)
            
            # Get other elements using matched patterns
            contact_info = self._extract_contact_with_format(doc, format)
            
            return {
                'sections': sections,
                'companies': companies,
                'contact': contact_info,
                'matched_format': format.name,
                'match_score': matched.get('match_score', 0)
            }
            
        except Exception as e:
            logger.error(f"Error processing with format: {str(e)}")
            raise

    def _process_without_format(self, doc: docx.Document) -> Dict[str, Any]:
        """
        Fallback processing without a matched format
        """
        # Implement fallback processing logic
        pass

    def _extract_sections_with_format(self, 
                                    doc: docx.Document, 
                                    format: ResumeFormat) -> Dict[str, List[str]]:
        """
        Extract sections using format patterns
        """
        sections = {}
        try:
            for section_name, pattern in format.section_patterns.items():
                section_content = []
                in_section = False
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    
                    # Check section start
                    if pattern['pattern'].upper() in text.upper():
                        in_section = True
                        continue
                    
                    # Check section end
                    if in_section and any(
                        other_pattern['pattern'].upper() in text.upper()
                        for other_name, other_pattern in format.section_patterns.items()
                        if other_name != section_name
                    ):
                        in_section = False
                        
                    # Collect section content
                    if in_section and text:
                        section_content.append(text)
                
                if section_content:
                    sections[section_name] = section_content
                    
            return sections
            
        except Exception as e:
            logger.error(f"Error extracting sections with format: {str(e)}")
            return {}

    def _extract_companies_with_format(self, 
                                     doc: docx.Document, 
                                     format: ResumeFormat) -> List[Dict[str, Any]]:
        """
        Extract company information using format patterns
        """
        companies = []
        try:
            if not format.company_patterns:
                return []
                
            pattern = format.company_patterns[0]  # Use first pattern as reference
            
            for para in doc.paragraphs:
                # Check if paragraph matches company pattern
                if (any(run.bold == pattern['formatting']['bold'] 
                       for run in para.runs if run.text.strip())):
                    
                    companies.append({
                        'name': para.text.strip(),
                        'location': len(companies),  # Track order
                        'text': para.text
                    })
            
            return companies
            
        except Exception as e:
            logger.error(f"Error extracting companies with format: {str(e)}")
            return []

    def _extract_contact_with_format(self, 
                                   doc: docx.Document, 
                                   format: ResumeFormat) -> Dict[str, str]:
        """
        Extract contact information using format patterns
        """
        contact = {}
        try:
            text = "\n".join(p.text for p in doc.paragraphs)
            
            # Extract email
            if format.email_pattern:
                import re
                email_match = re.search(format.email_pattern['regex'], text)
                if email_match:
                    contact['email'] = email_match.group()
            
            # Extract phone
            if format.phone_pattern:
                phone_match = re.search(format.phone_pattern['regex'], text)
                if phone_match:
                    contact['phone'] = phone_match.group()
            
            return contact
            
        except Exception as e:
            logger.error(f"Error extracting contact with format: {str(e)}")
            return {}