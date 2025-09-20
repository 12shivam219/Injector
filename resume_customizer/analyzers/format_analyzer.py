"""
Resume Format Analyzer Module
Analyzes and extracts format patterns from resumes
"""

import re
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import logging
from docx import Document
import hashlib

from infrastructure.utilities.logger import get_logger
from utilities.fuzzy_matcher import FuzzyMatcher
from database.models import BaseModel
from database.models import ResumeFormat, ResumeFormatMatch

logger = get_logger()

class FormatAnalyzer:
    """Analyzes resume formats and extracts patterns"""
    
    def __init__(self):
        self.logger = logger
        self.fuzzy_matcher = FuzzyMatcher()
        
        # Common regex patterns
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        self.phone_pattern = r'\b(\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b'
        self.date_pattern = r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b'
        
        # Section markers
        self.section_markers = [
            'EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS',
            'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE',
            'TECHNICAL SKILLS', 'CERTIFICATIONS'
        ]

    def save_format(self, session, name: str, description: str, patterns: Dict[str, Any]) -> str:
        """Save format patterns to database"""
        try:
            format_model = ResumeFormat(
                name=name,
                description=description,
                name_pattern=patterns.get('name_pattern'),
                email_pattern=patterns.get('email_pattern'),
                phone_pattern=patterns.get('phone_pattern'),
                section_patterns=patterns.get('section_patterns'),
                company_patterns=patterns.get('company_patterns'),
                title_patterns=patterns.get('title_patterns'),
                version='1.0'  # Set default version
            )
            
            session.add(format_model)
            session.commit()
            return str(format_model.id)
        except Exception as e:
            session.rollback()
            self.logger.error(f"Error saving format: {str(e)}")
            raise

    def analyze_format(self, doc: Document) -> Dict[str, Any]:
        """
        Analyze a resume document and extract its format patterns
        
        Args:
            doc: The Word document to analyze
            
        Returns:
            Dict containing format patterns
        """
        try:
            # Get document text
            full_text = self._get_document_text(doc)
            
            # Generate document hash
            doc_hash = self._generate_hash(full_text)
            
            # Extract patterns
            patterns = {
                'doc_hash': doc_hash,
                'name_pattern': self._extract_name_pattern(doc),
                'email_pattern': self._extract_contact_pattern(full_text, self.email_pattern),
                'phone_pattern': self._extract_contact_pattern(full_text, self.phone_pattern),
                'section_patterns': self._extract_section_patterns(doc),
                'company_patterns': self._extract_company_patterns(doc),
                'title_patterns': self._extract_title_patterns(doc)
            }
            
            return patterns
            
        except Exception as e:
            self.logger.error(f"Error analyzing format: {str(e)}")
            raise

    def _generate_hash(self, text: str) -> str:
        """Generate a hash of the document text"""
        return hashlib.sha256(text.encode()).hexdigest()

    def _get_document_text(self, doc: Document) -> str:
        """Extract full text from document"""
        return "\n".join(p.text for p in doc.paragraphs)

    def _extract_name_pattern(self, doc: Document) -> Dict[str, str]:
        """Extract the pattern for how names appear in the resume"""
        try:
            # Usually the name is in the first few paragraphs
            # Often larger font or bold
            def _run_font_size_gt(run, threshold=12):
                size = getattr(run.font, 'size', None)
                if size is None:
                    return False
                # size may be a Pt object with .pt attribute
                try:
                    value = getattr(size, 'pt', size)
                    return float(value) > threshold
                except Exception:
                    return False

            for i, para in enumerate(doc.paragraphs[:5]):
                if para.text.strip() and any((getattr(run, 'bold', False) is True) or _run_font_size_gt(run) for run in para.runs):
                    before_context = doc.paragraphs[i-1].text if i > 0 else ""
                    after_context = doc.paragraphs[i+1].text if i < len(doc.paragraphs)-1 else ""
                    return {
                        'start_marker': before_context[-50:] if before_context else "",
                        'end_marker': after_context[:50] if after_context else "",
                        'pattern': para.text.strip(),
                        'location': f"paragraph_{i+1}"
                    }
            return {}
        except Exception as e:
            self.logger.error(f"Error extracting name pattern: {str(e)}")
            return {}

    def _extract_contact_pattern(self, text: str, pattern: str) -> Dict[str, str]:
        """Extract pattern for contact information"""
        try:
            match = re.search(pattern, text)
            if match:
                start, end = match.span()
                return {
                    'start_marker': text[max(0, start-50):start],
                    'end_marker': text[end:min(len(text), end+50)],
                    'regex': pattern,
                    'sample': match.group()
                }
            return {}
        except Exception as e:
            self.logger.error(f"Error extracting contact pattern: {str(e)}")
            return {}

    def _extract_section_patterns(self, doc: Document) -> Dict[str, Dict[str, str]]:
        """Extract patterns for different resume sections"""
        sections = {}
        try:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip().upper()
                if any(marker in text for marker in self.section_markers):
                    # Found a section header
                    section_name = text
                    before_context = doc.paragraphs[i-1].text if i > 0 else ""
                    after_context = doc.paragraphs[i+1].text if i < len(doc.paragraphs)-1 else ""
                    
                    sections[section_name] = {
                        'start_marker': before_context[-50:] if before_context else "",
                        'end_marker': after_context[:50] if after_context else "",
                        'pattern': text,
                        'formatting': {
                            'bold': any(run.bold for run in para.runs),
                            'font_size': para.runs[0].font.size if para.runs else None
                        }
                    }
            return sections
        except Exception as e:
            self.logger.error(f"Error extracting section patterns: {str(e)}")
            return {}

    def _extract_company_patterns(self, doc: Document) -> List[Dict[str, str]]:
        """Extract patterns for how companies are listed"""
        companies = []
        try:
            in_experience = False
            current_company = {}
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                    
                # Check if we're in experience section
                if "EXPERIENCE" in text.upper():
                    in_experience = True
                    continue
                
                if in_experience:
                    # Look for company entries (often bold, may be followed by dates)
                    is_bold = any(getattr(run, 'bold', False) for run in para.runs)
                    has_date = bool(re.search(self.date_pattern, text))

                    # Accept bold entries as company names even if dates are not present
                    if is_bold:
                        font_size = None
                        try:
                            first_run = para.runs[0] if para.runs else None
                            size = getattr(first_run.font, 'size', None) if first_run is not None else None
                            font_size = getattr(size, 'pt', size) if size is not None else None
                        except Exception:
                            font_size = None

                        current_company = {
                            'pattern': text,
                            'start_marker': doc.paragraphs[i-1].text[-50:] if i > 0 else "",
                            'end_marker': doc.paragraphs[i+1].text[:50] if i < len(doc.paragraphs)-1 else "",
                            'formatting': {
                                'bold': True,
                                'font_size': font_size
                            }
                        }
                        companies.append(current_company)
            
            return companies
        except Exception as e:
            self.logger.error(f"Error extracting company patterns: {str(e)}")
            return []

    def _extract_title_patterns(self, doc: Document) -> List[Dict[str, str]]:
        """Extract patterns for job titles"""
        titles = []
        try:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Titles often appear after company names, are often bold or italicized
                if any(run.bold or run.italic for run in para.runs):
                    titles.append({
                        'pattern': text,
                        'start_marker': doc.paragraphs[i-1].text[-50:] if i > 0 else "",
                        'end_marker': doc.paragraphs[i+1].text[:50] if i < len(doc.paragraphs)-1 else "",
                        'formatting': {
                            'bold': any(run.bold for run in para.runs),
                            'italic': any(run.italic for run in para.runs),
                            'font_size': para.runs[0].font.size if para.runs else None
                        }
                    })
            
            return titles
        except Exception as e:
            self.logger.error(f"Error extracting title patterns: {str(e)}")
            return []

    def match_format(self, doc: Document, stored_format: ResumeFormat) -> Tuple[int, Dict[str, Any]]:
        """
        Try to match a document against a stored format
        
        Args:
            doc: Document to match
            stored_format: Format to match against
            
        Returns:
            Tuple of (match_score, matched_elements)
        """
        try:
            # Extract patterns from new document
            new_patterns = self.analyze_format(doc)
            
            # Initialize scoring
            total_score = 0
            matched_elements = {}
            
            # Compare patterns
            if self._match_contact_patterns(new_patterns, stored_format):
                total_score += 30
                matched_elements['contact'] = True
                
            section_score = self._match_sections(new_patterns, stored_format)
            total_score += section_score * 30
            matched_elements['sections'] = section_score
            
            company_score = self._match_companies(new_patterns, stored_format)
            total_score += company_score * 40
            matched_elements['companies'] = company_score
            
            return total_score, matched_elements
            
        except Exception as e:
            self.logger.error(f"Error matching format: {str(e)}")
            return 0, {}

    def _match_contact_patterns(self, new_patterns: Dict, stored_format: ResumeFormat) -> bool:
        """Check if contact information patterns match"""
        try:
            # Check email pattern
            if new_patterns.get('email_pattern') and stored_format.email_pattern:
                if new_patterns['email_pattern']['regex'] != stored_format.email_pattern['regex']:
                    return False
                    
            # Check phone pattern
            if new_patterns.get('phone_pattern') and stored_format.phone_pattern:
                if new_patterns['phone_pattern']['regex'] != stored_format.phone_pattern['regex']:
                    return False
                    
            return True
        except:
            return False

    def _match_sections(self, new_patterns: Dict, stored_format: ResumeFormat) -> float:
        """Calculate how well sections match"""
        try:
            # If the stored format does not define section patterns, consider sections as matched
            if not stored_format.section_patterns:
                return 1.0

            if not new_patterns.get('section_patterns'):
                return 0.0
                
            matches = 0
            total = len(stored_format.section_patterns)
            
            for section, pattern in stored_format.section_patterns.items():
                if section in new_patterns['section_patterns']:
                    matches += 1
                    
            return matches / total
        except:
            return 0

    def _match_companies(self, new_patterns: Dict, stored_format: ResumeFormat) -> float:
        """
        Calculate how well company patterns match using fuzzy string matching.
        Returns a score between 0 and 1 based on:
        - Formatting match (20% of score)
        - Company name fuzzy matches (80% of score)
        """
        try:
            if not new_patterns.get('company_patterns') or not stored_format.company_patterns:
                return 0
                
            # Match formatting (20% of total score)
            format_score = 0.0
            stored_format_style = stored_format.company_patterns[0]['formatting']
            new_format_style = new_patterns['company_patterns'][0]['formatting']
            
            if stored_format_style.get('bold') == new_format_style.get('bold'):
                format_score = 0.2
            
            # Match company names using fuzzy matching (80% of total score)
            name_matches = 0
            total_companies = len(stored_format.company_patterns)
            
            # For each stored company pattern, try to find a fuzzy match
            for stored_company in stored_format.company_patterns:
                stored_name = stored_company['pattern']
                
                # Try to find best match among new companies
                best_match_score = 0.0
                for new_company in new_patterns['company_patterns']:
                    new_name = new_company['pattern']
                    is_match, score = self.fuzzy_matcher.compare_strings(stored_name, new_name)
                    best_match_score = max(best_match_score, score / 100.0)  # Convert score to 0-1 range
                
                name_matches += best_match_score
            
            # Calculate final name matching score (0-0.8 range)
            name_score = 0.8 * (name_matches / total_companies if total_companies > 0 else 0)
            
            # Combine format and name scores
            return format_score + name_score
            
        except Exception as e:
            self.logger.error(f"Error in company matching: {str(e)}")
            return 0