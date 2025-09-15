"""
Bullet point formatter for document processing.
Handles detection and formatting of bullet points in Word documents.
"""
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
import logging
from infrastructure.utilities.logger import get_logger

from .base_formatters import DocumentFormatter, ListFormatterMixin

# Initialize logger with fallback
try:
    logger = get_logger()
except Exception:
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)


@dataclass
class BulletFormatting:
    """Data class to hold bullet formatting information."""
    runs_formatting: List[Dict[str, Any]]
    paragraph_formatting: Dict[str, Any]
    style: str
    bullet_marker: str
    bullet_separator: str
    list_format: Dict[str, Any]


class BulletFormatter(DocumentFormatter, ListFormatterMixin):
    """Handles bullet point formatting and style preservation."""
    
    def __init__(self, bullet_markers: List[str] = None):
        # Dash variants
        self.dash_variants = ['-', '–', '—']
        self.bullet_markers = bullet_markers or ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        
        self.bullet_patterns = [
            r'^\s*[-−–—]\s+',  # Various dash types
            r'^\s*[•·▪▫]\s+',  # Bullet symbols
            r'^\s*[*]\s+',     # Asterisk
            r'^\s*[+]\s+',     # Plus sign
        ]
        self.default_marker = '- '
    
    # -------------------------------
    # Formatting extraction
    # -------------------------------
    def extract_formatting(self, doc: DocumentType, paragraph_index: int) -> Optional[BulletFormatting]:
        try:
            if paragraph_index >= len(doc.paragraphs):
                return None
                
            para = doc.paragraphs[paragraph_index]
            if not self._is_bullet_point(para.text):
                return None
                
            formatting_info = BulletFormatting(
                runs_formatting=[],
                paragraph_formatting={},
                style=para.style.name if para.style else 'Normal',
                bullet_marker=self._extract_bullet_marker(para.text) or '-',
                bullet_separator=self._detect_bullet_separator(para.text) or ' ',
                list_format=self._extract_list_format(para)
            )
            
            # Extract run formatting
            for run in para.runs:
                try:
                    run_format = {
                        'text': run.text,
                        'font_name': run.font.name if hasattr(run.font, 'name') else None,
                        'font_size': run.font.size.pt if hasattr(run.font, 'size') and run.font.size else None,
                        'bold': run.font.bold if hasattr(run.font, 'bold') else None,
                        'italic': run.font.italic if hasattr(run.font, 'italic') else None,
                        'underline': run.font.underline if hasattr(run.font, 'underline') else None,
                        'color': run.font.color.rgb if (hasattr(run.font, 'color') and 
                                                     run.font.color and 
                                                     hasattr(run.font.color, 'rgb')) else None
                    }
                    formatting_info.runs_formatting.append(run_format)
                except Exception:
                    continue
            
            # Extract paragraph formatting
            if hasattr(para, 'paragraph_format'):
                p_format = para.paragraph_format
                formatting_info.paragraph_formatting = {
                    'alignment': p_format.alignment,
                    'first_line_indent': p_format.first_line_indent,
                    'left_indent': p_format.left_indent,
                    'right_indent': p_format.right_indent,
                    'space_before': p_format.space_before,
                    'space_after': p_format.space_after,
                    'line_spacing': p_format.line_spacing,
                    'keep_together': p_format.keep_together,
                    'keep_with_next': p_format.keep_with_next,
                    'page_break_before': p_format.page_break_before,
                    'widow_control': p_format.widow_control
                }
                
            return formatting_info
            
        except Exception:
            bullet_marker = '-'
            if 'para' in locals() and para.text:
                try:
                    bullet_marker = self._extract_bullet_marker(para.text)
                except Exception:
                    pass
            
            return BulletFormatting(
                runs_formatting=[{'text': para.text if 'para' in locals() else ''}],
                paragraph_formatting={},
                style='Normal',
                bullet_marker=bullet_marker,
                bullet_separator=' ',
                list_format={'ilvl': 0, 'numId': 1, 'style': 'List Bullet', 'indent': 0, 'is_list': False}
            )
    
    # -------------------------------
    # Apply formatting
    # -------------------------------
    def apply_formatting(self, paragraph: Paragraph, formatting: BulletFormatting, 
                        text: str, fallback_formatting: Optional[BulletFormatting] = None) -> None:
        """Apply extracted formatting to a new bullet point paragraph."""
        try:
            if not formatting and fallback_formatting:
                formatting = fallback_formatting

            clean_text = self._clean_bullet_text(text)
            bullet_marker = formatting.bullet_marker.strip() if formatting and formatting.bullet_marker else '-'
            bullet_separator = formatting.bullet_separator if formatting and formatting.bullet_separator else " "

            # ✅ Dash bullets → plain text only
            if bullet_marker in self.dash_variants:
                paragraph.clear()
                paragraph.add_run(f"{bullet_marker}{bullet_separator}{clean_text}")
                logger.debug(f"Applied dash bullet formatting with marker '{bullet_marker}'")
                return

            # ✅ Real bullets → apply Word list style
            if formatting and formatting.style:
                try:
                    paragraph.style = formatting.style
                except Exception:
                    pass

            if formatting and formatting.paragraph_formatting:
                pf = paragraph.paragraph_format
                for attr, value in formatting.paragraph_formatting.items():
                    if value is not None:
                        try:
                            setattr(pf, attr, value)
                        except Exception:
                            continue

            if formatting and formatting.list_format and formatting.list_format.get("is_list"):
                self._apply_list_formatting(paragraph, formatting.list_format)

            paragraph.clear()
            paragraph.add_run(clean_text)

        except Exception as e:
            logger.warning(f"Formatting application failed, using fallback: {e}")
            self._apply_basic_formatting(paragraph, text, formatting)
    
    def _apply_basic_formatting(self, paragraph: Paragraph, text: str, 
                              formatting: Optional[BulletFormatting] = None) -> None:
        """Apply basic bullet formatting as a fallback."""
        try:
            marker = formatting.bullet_marker.strip() if formatting and formatting.bullet_marker else '-'
            separator = formatting.bullet_separator if formatting and formatting.bullet_separator else ' '
            clean_text = self._clean_bullet_text(text)
            paragraph.clear()

            if marker in self.dash_variants:
                # ✅ Dash bullets → plain text only
                paragraph.add_run(f"{marker}{separator}{clean_text}")
                logger.debug(f"Applied basic dash bullet formatting with marker '{marker}'")
            else:
                # ✅ Real bullets → rely on Word list style, no manual marker
                paragraph.add_run(clean_text)
                logger.debug(f"Applied basic real bullet formatting with marker '{marker}'")

        except Exception as e:
            logger.error(f"Basic formatting failed: {e}")
            paragraph.clear()
            paragraph.add_run(text)
    
    # -------------------------------
    # Helpers
    # -------------------------------
    def _extract_list_format(self, paragraph: Paragraph) -> Dict[str, Any]:
        list_format = {
            'ilvl': 0,
            'numId': 1,
            'style': 'List Bullet',
            'indent': 0,
            'is_list': False
        }
        try:
            if hasattr(paragraph, 'style') and paragraph.style is not None:
                list_format['style'] = paragraph.style.name
            
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format is not None:
                if hasattr(paragraph.paragraph_format, 'left'):
                    list_format['indent'] = paragraph.paragraph_format.left or 0
            
            if not self._is_bullet_point(paragraph.text):
                return list_format
                
            list_format['is_list'] = True
            
            if not hasattr(paragraph, '_element') or not hasattr(paragraph._element, 'pPr'):
                return list_format
                
            pPr = paragraph._element.pPr
            if pPr is None:
                return list_format
            
            numPr = getattr(pPr, 'numPr', None)
            if numPr is None:
                return list_format
            
            if hasattr(numPr, 'ilvl') and hasattr(numPr.ilvl, 'val'):
                list_format['ilvl'] = numPr.ilvl.val
            
            if hasattr(numPr, 'numId') and hasattr(numPr.numId, 'val'):
                list_format['numId'] = numPr.numId.val
                        
        except Exception:
            pass
            
        return list_format
    
    def _is_bullet_point(self, text: str) -> bool:
        text = text.strip()
        return any(
            text.startswith(marker + ' ') or text.startswith(marker)
            for marker in self.bullet_markers
        ) or (text and text[0].isdigit() and '.' in text[:3])
    
    def detect_document_bullet_marker(self, document: DocumentType) -> str:
        import re
        marker_counts = {}
        bullet_point_count = 0
        
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            if any(keyword in text.lower() for keyword in ['experience', 'education', 'skills', 'summary', 'objective']):
                continue
            
            for pattern in self.bullet_patterns:
                match = re.match(pattern, text)
                if match:
                    bullet_point_count += 1
                    marker = match.group().strip().rstrip(' \t')
                    marker_counts[marker] = marker_counts.get(marker, 0) + 1
                    break
            
            for marker in self.bullet_markers:
                if text.startswith(marker + ' ') or text.startswith(marker + '\t'):
                    bullet_point_count += 1
                    marker_counts[marker] = marker_counts.get(marker, 0) + 1
                    break
        
        if marker_counts:
            return max(marker_counts, key=marker_counts.get)
        return self.default_marker
    
    def _extract_bullet_marker(self, text: str) -> str:
        import re
        text = text.strip()
        
        for pattern in self.bullet_patterns:
            match = re.match(pattern, text)
            if match:
                return match.group().strip().rstrip(' \t') or '-'
        
        all_markers = ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        for marker in all_markers:
            if text.startswith(marker + '\t') or text.startswith(marker + ' '):
                return marker
            elif text.startswith(marker) and len(text) > 1 and not text[1].isalnum():
                return marker
                
        if text and text[0].isdigit():
            for i, char in enumerate(text):
                if char in '.)': 
                    return text[:i+1]
        
        return '-'
    
    def _detect_bullet_separator(self, text: str) -> str:
        text = text.strip()
        markers = ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        
        for marker in markers:
            if text.startswith(marker + '\t'):
                return '\t'
            elif text.startswith(marker + ' '):
                return ' '
        
        return '\t'
    
    def _clean_bullet_text(self, text: str) -> str:
        dash_and_bullet_chars = '-–—•●*◦▪▫‣ \t'
        clean_text = text.lstrip(dash_and_bullet_chars).lstrip()
        
        if clean_text and clean_text[0].isdigit():
            for i, char in enumerate(clean_text):
                if char in '.)':
                    clean_text = clean_text[i+1:].lstrip()
                    break
        
        return clean_text
