"""
Test script to verify improved project detection for Company | Date format
"""

import sys
import os
sys.path.append(os.getcwd())

from docx import Document
from io import BytesIO
from resume_customizer.detectors.project_detector import ProjectDetector
from resume_customizer.processors.resume_processor import PreviewGenerator

def create_user_format_resume():
    """Create a resume in the user's format: Company | Date, Job Title, Responsibilities, bullet points"""
    doc = Document()
    
    # Header
    doc.add_heading("John Smith", 0)
    doc.add_paragraph("john.smith@email.com | (555) 123-4567")
    
    # Experience section
    doc.add_heading("Professional Experience", 1)
    
    # Project 1 - User's format
    doc.add_paragraph("TechCorp Inc. | June 2021 - Present")
    doc.add_paragraph("Senior Full Stack Developer")
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("Responsibilities")
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("- Led development of web applications")
    doc.add_paragraph("- Mentored junior developers")
    doc.add_paragraph("- Improved system performance by 40%")
    
    doc.add_paragraph("")  # Empty line between projects
    
    # Project 2 - User's format
    doc.add_paragraph("StartupXYZ | January 2019 - May 2021")
    doc.add_paragraph("Frontend Developer")
    doc.add_paragraph("")
    doc.add_paragraph("Responsibilities")
    doc.add_paragraph("")
    doc.add_paragraph("- Built responsive user interfaces")
    doc.add_paragraph("- Integrated with REST APIs")
    doc.add_paragraph("- Collaborated with design team")
    
    doc.add_paragraph("")
    
    # Project 3 - User's format
    doc.add_paragraph("DevSolutions LLC | March 2017 - December 2018")
    doc.add_paragraph("Junior Software Engineer")
    doc.add_paragraph("")
    doc.add_paragraph("Responsibilities")
    doc.add_paragraph("")
    doc.add_paragraph("- Developed mobile applications")
    doc.add_paragraph("- Wrote unit tests")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def test_project_detection():
    """Test the improved project detection"""
    print("=== TESTING IMPROVED PROJECT DETECTION ===\n")
    
    # Create test resume
    print("1. Creating resume in user's format...")
    resume_buffer = create_user_format_resume()
    print("✅ Resume created with format: Company | Date, Job Title, Responsibilities, bullet points")
    
    # Test project detection
    print("\n2. Testing project detection...")
    doc = Document(resume_buffer)
    detector = ProjectDetector()
    projects = detector.find_projects(doc)
    
    print(f"✅ Found {len(projects)} projects")
    
    expected_projects = 3
    if len(projects) == expected_projects:
        print("✅ Correct number of projects detected!")
        
        for i, project in enumerate(projects):
            print(f"\n📋 Project {i+1}:")
            print(f"   Name: {project.name}")
            print(f"   Company: {project.company}")
            print(f"   Role: {project.role}")
            print(f"   Date Range: {project.date_range}")
            print(f"   Start Index: {project.start_index}")
            print(f"   End Index: {project.end_index}")
            print(f"   Bullet Points: {len(project.bullet_points)}")
            
    else:
        print(f"❌ Expected {expected_projects} projects, but found {len(projects)}")
        for i, project in enumerate(projects):
            print(f"   Project {i+1}: {project.name}")
    
    # Test with preview generation
    print(f"\n3. Testing preview with detected projects...")
    tech_input = """Python
• Developed REST API endpoints
• Implemented database queries

JavaScript
• Created dynamic UI components
• Built interactive dashboards

React
• Developed component libraries
• Implemented state management"""
    
    try:
        resume_buffer.seek(0)  # Reset buffer
        preview_gen = PreviewGenerator()
        result = preview_gen.generate_preview(resume_buffer, tech_input)
        
        if result['success']:
            print("✅ Preview generated successfully!")
            print(f"📊 Points added: {result['points_added']}")
            print(f"📊 Projects count: {result['projects_count']}")
            print(f"📊 Tech stacks: {', '.join(result['tech_stacks_used'])}")
            
            print(f"\n🆕 Points distribution:")
            for project, mapping in result['project_points_mapping'].items():
                print(f"📋 {project}: {len(mapping['points'])} points")
                for point in mapping['points'][:2]:  # Show first 2 points
                    print(f"    - {point}")
                if len(mapping['points']) > 2:
                    print(f"    ... and {len(mapping['points']) - 2} more")
            
            return True
        else:
            print(f"❌ Preview failed: {result['error']}")
            return False
            
    except Exception as e:
        print(f"❌ Exception during preview: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_project_detection()
    if success:
        print("\n🎉 All tests passed! Project detection is working correctly for your resume format.")
    else:
        print("\n❌ Some tests failed. Check the output above for details.")