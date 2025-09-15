"""
Test script to validate project detection with user's actual resume file
"""

import sys
import os
sys.path.append(os.getcwd())

from docx import Document
from io import BytesIO
from resume_customizer.detectors.project_detector import ProjectDetector
from resume_customizer.processors.resume_processor import PreviewGenerator

def test_actual_resume():
    """Test the improved project detection with the user's actual resume"""
    print("=== TESTING WITH ACTUAL RESUME ===\n")
    
    resume_path = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    # Check if file exists
    if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return False
    
    print(f"📄 Found resume file: {os.path.basename(resume_path)}")
    
    try:
        # Load the actual resume
        print("\n1. Loading actual resume...")
        doc = Document(resume_path)
        print("✅ Resume loaded successfully")
        
        # Show first few paragraphs for context
        print("\n📋 Resume structure (first 15 lines):")
        for i, para in enumerate(doc.paragraphs[:15]):
            text = para.text.strip()
            if text:
                print(f"   {i:2d}: {text}")
        
        # Test project detection
        print(f"\n2. Testing project detection...")
        detector = ProjectDetector()
        projects = detector.find_projects(doc)
        
        print(f"✅ Found {len(projects)} projects")
        
        if projects:
            print("\n📋 Detected Projects:")
            for i, project in enumerate(projects):
                print(f"\n   Project {i+1}:")
                print(f"      Name: {project.name}")
                print(f"      Company: {project.company}")
                print(f"      Role: {project.role}")
                print(f"      Date Range: {project.date_range}")
                print(f"      Start Index: {project.start_index}")
                print(f"      End Index: {project.end_index}")
                print(f"      Bullet Points: {len(project.bullet_points)}")
                
                # Show first few bullet points
                if project.bullet_points:
                    print(f"      Sample bullets:")
                    for bullet in project.bullet_points[:3]:
                        print(f"         - {bullet}")
                    if len(project.bullet_points) > 3:
                        print(f"         ... and {len(project.bullet_points) - 3} more")
        else:
            print("❌ No projects detected")
            return False
        
        # Test with preview generation
        print(f"\n3. Testing preview generation...")
        
        # Sample tech input based on common technologies
        tech_input = """Python
• Developed REST API endpoints using FastAPI
• Implemented data processing pipelines

JavaScript
• Created interactive web applications
• Built real-time data visualization dashboards

React
• Developed responsive user interface components
• Implemented state management with Redux

SQL
• Optimized database queries for performance
• Designed relational database schemas"""
        
        try:
            # Load file into BytesIO for preview generator
            with open(resume_path, 'rb') as file:
                file_buffer = BytesIO(file.read())
            
            preview_gen = PreviewGenerator()
            result = preview_gen.generate_preview(file_buffer, tech_input)
            
            if result['success']:
                print("✅ Preview generated successfully!")
                print(f"📊 Points added: {result['points_added']}")
                print(f"📊 Projects count: {result['projects_count']}")
                print(f"📊 Tech stacks: {', '.join(result['tech_stacks_used'])}")
                
                print(f"\n🆕 Points distribution:")
                for project, mapping in result['project_points_mapping'].items():
                    print(f"📋 {project}: {len(mapping['points'])} points")
                    for point in mapping['points'][:2]:  # Show first 2 points
                        print(f"      - {point}")
                    if len(mapping['points']) > 2:
                        print(f"      ... and {len(mapping['points']) - 2} more")
                
                # Optionally save preview to see the result
                if 'preview_doc' in result:
                    preview_path = r"C:\Users\HP\Downloads\RSInjector\preview_test_output.docx"
                    result['preview_doc'].save(preview_path)
                    print(f"\n💾 Preview saved to: {preview_path}")
                
                return True
            else:
                print(f"❌ Preview failed: {result['error']}")
                if 'debug_info' in result:
                    print(f"🐛 Debug info: {result['debug_info']}")
                return False
                
        except Exception as e:
            print(f"❌ Exception during preview: {e}")
            import traceback
            traceback.print_exc()
            return False
            
    except Exception as e:
        print(f"❌ Error loading resume: {e}")
        import traceback
        traceback.print_exc()
        return False

def analyze_resume_structure():
    """Analyze the structure of the resume to understand its format"""
    print("\n=== ANALYZING RESUME STRUCTURE ===\n")
    
    resume_path = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return
    
    try:
        doc = Document(resume_path)
        print(f"📄 Analyzing: {os.path.basename(resume_path)}")
        print(f"📊 Total paragraphs: {len(doc.paragraphs)}")
        
        print(f"\n📋 Full document structure:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Check formatting
                style_name = para.style.name if para.style else "Normal"
                font_size = "Unknown"
                is_bold = False
                
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.font.size:
                        font_size = f"{first_run.font.size.pt}pt"
                    if first_run.font.bold:
                        is_bold = True
                
                formatting = f"[{style_name}"
                if is_bold:
                    formatting += ", Bold"
                formatting += f", {font_size}]"
                
                print(f"   {i:2d}: {formatting:25} {text}")
        
        print(f"\n🔍 Looking for patterns:")
        detector = ProjectDetector()
        
        company_date_lines = []
        job_title_lines = []
        responsibility_lines = []
        bullet_lines = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            if detector._looks_like_company_date(text):
                company_date_lines.append((i, text))
            elif detector._is_responsibilities_heading(text):
                responsibility_lines.append((i, text))
            elif detector._is_bullet_point(text):
                bullet_lines.append((i, text))
            elif any(keyword in text.lower() for keyword in detector.config["job_title_keywords"]):
                job_title_lines.append((i, text))
        
        print(f"   📌 Company | Date lines: {len(company_date_lines)}")
        for i, text in company_date_lines:
            print(f"      Line {i}: {text}")
            
        print(f"   👔 Potential job title lines: {len(job_title_lines)}")
        for i, text in job_title_lines:
            print(f"      Line {i}: {text}")
            
        print(f"   📝 Responsibilities headers: {len(responsibility_lines)}")
        for i, text in responsibility_lines:
            print(f"      Line {i}: {text}")
            
        print(f"   • Bullet points: {len(bullet_lines)}")
        for i, text in bullet_lines[:5]:  # Show first 5
            print(f"      Line {i}: {text}")
        if len(bullet_lines) > 5:
            print(f"      ... and {len(bullet_lines) - 5} more bullets")
        
    except Exception as e:
        print(f"❌ Error analyzing resume: {e}")

if __name__ == "__main__":
    # First analyze the structure
    analyze_resume_structure()
    
    print("\n" + "="*60 + "\n")
    
    # Then test project detection
    success = test_actual_resume()
    
    if success:
        print("\n🎉 All tests passed! Your resume format is compatible with the improved project detection.")
    else:
        print("\n❌ Some tests failed. The resume format may need additional adjustments.")
        print("💡 Check the structure analysis above to understand the format better.")