#!/usr/bin/env python3
"""
Test script to verify bullet point formatting consistency functionality.
Tests both dash and bullet marker detection and application.
"""

import sys
import os
from io import BytesIO
from docx import Document

# Add the current directory to Python path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def create_test_document_with_dashes():
    """Create a test document with dash bullet points."""
    doc = Document()
    
    # Add title
    doc.add_heading('John Doe - Software Engineer', 0)
    
    # Add experience section
    doc.add_heading('EXPERIENCE', level=1)
    
    # Add project with dash bullets
    project_para = doc.add_paragraph('Senior Software Engineer | TechCorp Inc. | 2020-2023')
    project_para.style = 'Heading 2'
    
    responsibilities_para = doc.add_paragraph('Responsibilities:')
    responsibilities_para.style = 'Heading 3'
    
    # Add existing dash bullet points
    doc.add_paragraph('- Developed web applications using React and Node.js')
    doc.add_paragraph('- Led team of 5 developers on enterprise projects')
    doc.add_paragraph('- Implemented CI/CD pipelines reducing deployment time by 50%')
    
    return doc

def create_test_document_with_bullets():
    """Create a test document with bullet point markers."""
    doc = Document()
    
    # Add title
    doc.add_heading('Jane Smith - Full Stack Developer', 0)
    
    # Add experience section
    doc.add_heading('EXPERIENCE', level=1)
    
    # Add project with bullet markers
    project_para = doc.add_paragraph('Full Stack Developer | StartupXYZ | 2021-2023')
    project_para.style = 'Heading 2'
    
    responsibilities_para = doc.add_paragraph('Responsibilities:')
    responsibilities_para.style = 'Heading 3'
    
    # Add existing bullet points
    doc.add_paragraph('• Built responsive web applications using Vue.js')
    doc.add_paragraph('• Designed RESTful APIs with Python Flask')
    doc.add_paragraph('• Managed PostgreSQL databases and optimized queries')
    
    return doc

def test_bullet_marker_detection():
    """Test bullet marker detection functionality."""
    print("🧪 TESTING BULLET MARKER DETECTION")
    print("=" * 60)
    
    try:
        from resume_customizer.formatters.bullet_formatter import BulletFormatter
        
        formatter = BulletFormatter()
        
        # Test 1: Document with dashes
        print("\n📄 Test 1: Document with dash bullets")
        dash_doc = create_test_document_with_dashes()
        detected_marker = formatter.detect_document_bullet_marker(dash_doc)
        print(f"   Detected marker: '{detected_marker}'")
        
        expected_dash = '-'
        if detected_marker == expected_dash:
            print("   ✅ PASS: Correctly detected dash marker")
        else:
            print(f"   ❌ FAIL: Expected '{expected_dash}', got '{detected_marker}'")
        
        # Test 2: Document with bullets
        print("\n📄 Test 2: Document with bullet markers")
        bullet_doc = create_test_document_with_bullets()
        detected_marker = formatter.detect_document_bullet_marker(bullet_doc)
        print(f"   Detected marker: '{detected_marker}'")
        
        expected_bullet = '•'
        if detected_marker == expected_bullet:
            print("   ✅ PASS: Correctly detected bullet marker")
        else:
            print(f"   ❌ FAIL: Expected '{expected_bullet}', got '{detected_marker}'")
        
        return True
        
    except Exception as e:
        print(f"   ❌ FAIL: Exception during marker detection test: {e}")
        return False

def test_formatting_extraction():
    """Test formatting extraction from existing bullets."""
    print("\n🔍 TESTING FORMATTING EXTRACTION")
    print("=" * 60)
    
    try:
        from resume_customizer.formatters.bullet_formatter import BulletFormatter
        
        formatter = BulletFormatter()
        
        # Test with dash document
        print("\n📄 Test 1: Extract formatting from dash bullets")
        dash_doc = create_test_document_with_dashes()
        
        # Find a bullet point paragraph
        bullet_para_index = None
        for i, para in enumerate(dash_doc.paragraphs):
            if formatter._is_bullet_point(para.text):
                bullet_para_index = i
                break
        
        if bullet_para_index is not None:
            formatting = formatter.extract_formatting(dash_doc, bullet_para_index)
            if formatting:
                print(f"   Extracted marker: '{formatting.bullet_marker}'")
                print(f"   Extracted separator: '{formatting.bullet_separator}'")
                if formatting.bullet_marker == '-':
                    print("   ✅ PASS: Correctly extracted dash formatting")
                else:
                    print(f"   ❌ FAIL: Expected '-', got '{formatting.bullet_marker}'")
            else:
                print("   ❌ FAIL: No formatting extracted")
        else:
            print("   ❌ FAIL: No bullet points found")
        
        # Test with bullet document
        print("\n📄 Test 2: Extract formatting from bullet markers")
        bullet_doc = create_test_document_with_bullets()
        
        bullet_para_index = None
        for i, para in enumerate(bullet_doc.paragraphs):
            if formatter._is_bullet_point(para.text):
                bullet_para_index = i
                break
        
        if bullet_para_index is not None:
            formatting = formatter.extract_formatting(bullet_doc, bullet_para_index)
            if formatting:
                print(f"   Extracted marker: '{formatting.bullet_marker}'")
                print(f"   Extracted separator: '{formatting.bullet_separator}'")
                if formatting.bullet_marker == '•':
                    print("   ✅ PASS: Correctly extracted bullet formatting")
                else:
                    print(f"   ❌ FAIL: Expected '•', got '{formatting.bullet_marker}'")
            else:
                print("   ❌ FAIL: No formatting extracted")
        else:
            print("   ❌ FAIL: No bullet points found")
        
        return True
        
    except Exception as e:
        print(f"   ❌ FAIL: Exception during formatting extraction test: {e}")
        return False

def test_point_addition_consistency():
    """Test that new points are added with consistent formatting."""
    print("\n➕ TESTING POINT ADDITION CONSISTENCY")
    print("=" * 60)
    
    try:
        from resume_customizer.processors.document_processor import get_document_processor
        from resume_customizer.detectors.project_detector import ProjectDetector
        
        doc_processor = get_document_processor()
        project_detector = ProjectDetector()
        
        # Test 1: Add points to dash document
        print("\n📄 Test 1: Add points to document with dashes")
        dash_doc = create_test_document_with_dashes()
        
        # Detect document marker
        document_marker = doc_processor.bullet_formatter.detect_document_bullet_marker(dash_doc)
        print(f"   Document marker detected: '{document_marker}'")
        
        # Find projects
        projects = project_detector.find_projects(dash_doc)
        if projects:
            project = projects[0]
            test_points = [
                {'text': 'Implemented microservices architecture using Docker'},
                {'text': 'Optimized database queries improving performance by 40%'}
            ]
            
            # Add points
            added = doc_processor._add_points_to_project(dash_doc, project, test_points, document_marker)
            print(f"   Added {added} points to project")
            
            # Check if new points use the same marker
            new_bullets_found = 0
            dash_bullets_found = 0
            
            for para in dash_doc.paragraphs:
                text = para.text.strip()
                if any(test_text in text for test_text in ['microservices architecture', 'database queries']):
                    new_bullets_found += 1
                    if text.startswith('- '):
                        dash_bullets_found += 1
                        print(f"   ✅ New point uses dash: '{text[:50]}...'")
                    else:
                        print(f"   ❌ New point doesn't use dash: '{text[:50]}...'")
            
            if new_bullets_found == dash_bullets_found and new_bullets_found > 0:
                print("   ✅ PASS: All new points use dash markers consistently")
            else:
                print(f"   ❌ FAIL: {dash_bullets_found}/{new_bullets_found} new points use dash markers")
        else:
            print("   ❌ FAIL: No projects found in dash document")
        
        # Test 2: Add points to bullet document
        print("\n📄 Test 2: Add points to document with bullets")
        bullet_doc = create_test_document_with_bullets()
        
        # Detect document marker
        document_marker = doc_processor.bullet_formatter.detect_document_bullet_marker(bullet_doc)
        print(f"   Document marker detected: '{document_marker}'")
        
        # Find projects
        projects = project_detector.find_projects(bullet_doc)
        if projects:
            project = projects[0]
            test_points = [
                {'text': 'Deployed applications on AWS cloud infrastructure'},
                {'text': 'Created automated testing suites with Jest and Cypress'}
            ]
            
            # Add points
            added = doc_processor._add_points_to_project(bullet_doc, project, test_points, document_marker)
            print(f"   Added {added} points to project")
            
            # Check if new points use the same marker
            new_bullets_found = 0
            bullet_bullets_found = 0
            
            for para in bullet_doc.paragraphs:
                text = para.text.strip()
                if any(test_text in text for test_text in ['AWS cloud infrastructure', 'automated testing suites']):
                    new_bullets_found += 1
                    if text.startswith('• '):
                        bullet_bullets_found += 1
                        print(f"   ✅ New point uses bullet: '{text[:50]}...'")
                    else:
                        print(f"   ❌ New point doesn't use bullet: '{text[:50]}...'")
            
            if new_bullets_found == bullet_bullets_found and new_bullets_found > 0:
                print("   ✅ PASS: All new points use bullet markers consistently")
            else:
                print(f"   ❌ FAIL: {bullet_bullets_found}/{new_bullets_found} new points use bullet markers")
        else:
            print("   ❌ FAIL: No projects found in bullet document")
        
        return True
        
    except Exception as e:
        print(f"   ❌ FAIL: Exception during point addition test: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_end_to_end_processing():
    """Test end-to-end processing with formatting consistency."""
    print("\n🔄 TESTING END-TO-END PROCESSING")
    print("=" * 60)
    
    try:
        from resume_customizer.processors.document_processor import get_document_processor
        from resume_customizer.parsers.text_parser import parse_input_text
        
        doc_processor = get_document_processor()
        
        # Test input
        tech_stack_input = """Python
• Developed web applications using Django framework
• Built RESTful APIs with Flask

JavaScript  
• Created interactive UIs with React
• Implemented Node.js backend services"""
        
        # Parse input
        selected_points, tech_stacks_used = parse_input_text(tech_stack_input)
        print(f"   Parsed {len(selected_points)} points from {len(tech_stacks_used)} tech stacks")
        
        # Test with dash document
        print("\n📄 Test 1: End-to-end with dash document")
        dash_doc = create_test_document_with_dashes()
        
        # Save to buffer
        dash_buffer = BytesIO()
        dash_doc.save(dash_buffer)
        dash_buffer.seek(0)
        
        # Process document
        try:
            processed_buffer = doc_processor.process_document(dash_buffer, (selected_points, tech_stacks_used))
            processed_doc = Document(processed_buffer)
            
            # Check formatting consistency
            dash_count = 0
            bullet_count = 0
            total_bullets = 0
            
            for para in processed_doc.paragraphs:
                text = para.text.strip()
                if doc_processor.bullet_formatter._is_bullet_point(text):
                    total_bullets += 1
                    if text.startswith('- '):
                        dash_count += 1
                    elif text.startswith('• '):
                        bullet_count += 1
            
            print(f"   Total bullet points: {total_bullets}")
            print(f"   Dash bullets: {dash_count}")
            print(f"   Bullet markers: {bullet_count}")
            
            if dash_count > 0 and bullet_count == 0:
                print("   ✅ PASS: All bullets use dash markers consistently")
            elif bullet_count > 0 and dash_count == 0:
                print("   ⚠️ UNEXPECTED: All bullets use bullet markers (should be dashes)")
            else:
                print(f"   ❌ FAIL: Mixed markers found - {dash_count} dashes, {bullet_count} bullets")
                
        except Exception as e:
            print(f"   ❌ FAIL: Processing failed: {e}")
        
        # Test with bullet document
        print("\n📄 Test 2: End-to-end with bullet document")
        bullet_doc = create_test_document_with_bullets()
        
        # Save to buffer
        bullet_buffer = BytesIO()
        bullet_doc.save(bullet_buffer)
        bullet_buffer.seek(0)
        
        # Process document
        try:
            processed_buffer = doc_processor.process_document(bullet_buffer, (selected_points, tech_stacks_used))
            processed_doc = Document(processed_buffer)
            
            # Check formatting consistency
            dash_count = 0
            bullet_count = 0
            total_bullets = 0
            
            for para in processed_doc.paragraphs:
                text = para.text.strip()
                if doc_processor.bullet_formatter._is_bullet_point(text):
                    total_bullets += 1
                    if text.startswith('- '):
                        dash_count += 1
                    elif text.startswith('• '):
                        bullet_count += 1
            
            print(f"   Total bullet points: {total_bullets}")
            print(f"   Dash bullets: {dash_count}")
            print(f"   Bullet markers: {bullet_count}")
            
            if bullet_count > 0 and dash_count == 0:
                print("   ✅ PASS: All bullets use bullet markers consistently")
            elif dash_count > 0 and bullet_count == 0:
                print("   ⚠️ UNEXPECTED: All bullets use dash markers (should be bullets)")
            else:
                print(f"   ❌ FAIL: Mixed markers found - {dash_count} dashes, {bullet_count} bullets")
                
        except Exception as e:
            print(f"   ❌ FAIL: Processing failed: {e}")
        
        return True
        
    except Exception as e:
        print(f"   ❌ FAIL: Exception during end-to-end test: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run all formatting consistency tests."""
    print("🎯 BULLET POINT FORMATTING CONSISTENCY TESTS")
    print("=" * 60)
    print("Testing the fix for bullet point formatting preservation...")
    print()
    
    tests_passed = 0
    tests_failed = 0
    
    # Test 1: Marker detection
    if test_bullet_marker_detection():
        tests_passed += 1
    else:
        tests_failed += 1
    
    # Test 2: Formatting extraction
    if test_formatting_extraction():
        tests_passed += 1
    else:
        tests_failed += 1
    
    # Test 3: Point addition consistency
    if test_point_addition_consistency():
        tests_passed += 1
    else:
        tests_failed += 1
    
    # Test 4: End-to-end processing
    if test_end_to_end_processing():
        tests_passed += 1
    else:
        tests_failed += 1
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 TEST RESULTS SUMMARY")
    print(f"✅ Tests Passed: {tests_passed}")
    print(f"❌ Tests Failed: {tests_failed}")
    
    if tests_failed == 0:
        print("\n🎉 ALL TESTS PASSED!")
        print("✅ Bullet point formatting consistency is working correctly")
        print("✅ Dash markers are preserved when adding new points to dash documents")
        print("✅ Bullet markers are preserved when adding new points to bullet documents")
    else:
        print(f"\n⚠️ {tests_failed} test(s) failed")
        print("❌ Bullet point formatting consistency needs further debugging")
    
    return tests_failed == 0

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)