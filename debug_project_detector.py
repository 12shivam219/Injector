#!/usr/bin/env python3
"""
Debug script to trace through the ProjectDetector logic step by step
and understand why no projects are being detected.
"""

import sys
import os
from docx import Document

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def debug_project_detector():
    """Debug the ProjectDetector step by step."""
    print("🔍 Debugging ProjectDetector Logic")
    print("=" * 50)
    
    try:
        from resume_customizer.detectors.project_detector import ProjectDetector
        
        # Create a detector with debug info
        detector = ProjectDetector()
        
        resume_file = "Resume Format 1.docx"
        if not os.path.exists(resume_file):
            print(f"❌ Missing {resume_file}")
            return
            
        doc = Document(resume_file)
        
        # Manually trace through the find_projects logic
        projects = []
        current_project = None
        in_responsibilities = False
        in_experience_section = False
        i = 0
        
        print(f"Starting to process {len(doc.paragraphs)} paragraphs...")
        
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                print(f"Para {i:2d}: (empty) - skipping")
                i += 1
                continue
            
            print(f"Para {i:2d}: '{text}'")
            
            # Check if this is a section heading (like Professional Experience)
            is_section_heading = detector._is_section_heading(text)
            if is_section_heading:
                in_experience_section = True
                print(f"    ✅ SECTION HEADING detected! in_experience_section = True")
                i += 1
                continue
            
            print(f"    📍 in_experience_section = {in_experience_section}")
            
            # Only look for projects if we're in an experience section
            if not in_experience_section:
                print(f"    ⏭️  Not in experience section, skipping")
                i += 1
                continue
            
            # Check if this is a Company | Date format (highest priority)
            is_company_date = detector._looks_like_company_date(text)
            print(f"    🏢 _looks_like_company_date: {is_company_date}")
            
            if is_company_date:
                # Save previous project if exists
                if current_project:
                    projects.append(current_project)
                    print(f"    💾 Saved previous project: {current_project.name}")
                
                # Parse the company and date
                role, company, date_range = detector._parse_project_header(text)
                print(f"    📝 Parsed: role='{role}', company='{company}', date_range='{date_range}'")
                
                current_project = {
                    'name': role or company or f"Project {len(projects) + 1}",
                    'start_index': i,
                    'end_index': i,
                    'role': role,
                    'company': company,
                    'date_range': date_range,
                    'bullet_points': []
                }
                print(f"    🆕 Created new project: {current_project['name']}")
                in_responsibilities = False
                
            # Check if this could be a standalone job title (only if no current project)
            is_potential_project = detector._is_potential_project(text) and not current_project
            print(f"    🎯 _is_potential_project (and no current project): {is_potential_project}")
            
            if is_potential_project:
                # Save previous project if exists
                if current_project:
                    projects.append(current_project)
                    print(f"    💾 Saved previous project: {current_project['name']}")
                
                # Start new project
                role, company, date_range = detector._parse_project_header(text)
                print(f"    📝 Parsed: role='{role}', company='{company}', date_range='{date_range}'")
                
                current_project = {
                    'name': role or company or f"Project {len(projects) + 1}",
                    'start_index': i,
                    'end_index': i,
                    'role': role,
                    'company': company,
                    'date_range': date_range,
                    'bullet_points': []
                }
                print(f"    🆕 Created new project: {current_project['name']}")
                in_responsibilities = False
            
            # Check if we're in a responsibilities section
            is_responsibilities = detector._is_responsibilities_heading(text)
            print(f"    📋 _is_responsibilities_heading: {is_responsibilities}")
            
            if is_responsibilities:
                in_responsibilities = True
                print(f"    ✅ RESPONSIBILITIES detected! in_responsibilities = True")
                
            # If we have an active project, collect bullet points
            is_bullet = detector._is_bullet_point(text)
            print(f"    🔸 _is_bullet_point: {is_bullet}")
            print(f"    📍 current_project exists: {current_project is not None}")
            print(f"    📍 in_responsibilities: {in_responsibilities}")
            
            if current_project and (in_responsibilities or is_bullet):
                current_project['bullet_points'].append(text)
                current_project['end_index'] = i
                print(f"    ✅ Added bullet point to project: '{text}'")
            
            print()  # Empty line for readability
            i += 1
        
        # Add the last project if exists
        if current_project:
            projects.append(current_project)
            print(f"💾 Added final project: {current_project['name']}")
            
        print("=" * 50)
        print(f"🎯 FINAL RESULT: Found {len(projects)} projects")
        for j, project in enumerate(projects):
            print(f"  Project {j+1}: {project['name']} ({len(project['bullet_points'])} bullet points)")
        print("=" * 50)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_project_detector()