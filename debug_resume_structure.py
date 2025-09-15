#!/usr/bin/env python3
"""
Debug script to examine the structure of the created resume files
and understand why project detection is failing.
"""

import sys
import os
from docx import Document

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def debug_resume_structure():
    """Debug the structure of the resume files."""
    print("🔍 Debugging Resume File Structure")
    print("=" * 50)
    
    resume_files = ["Resume Format 1.docx", "Resume Format 3.docx"]
    
    for resume_file in resume_files:
        if not os.path.exists(resume_file):
            print(f"❌ Missing {resume_file}")
            continue
            
        print(f"\n📄 Analyzing {resume_file}")
        print("-" * 30)
        
        try:
            doc = Document(resume_file)
            
            print(f"Total paragraphs: {len(doc.paragraphs)}")
            print()
            
            # Check each paragraph
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:  # Only show non-empty paragraphs
                    print(f"Para {i:2d}: '{text}'")
                    
                    # Check specific conditions
                    if text.lower().startswith('responsibilities'):
                        print(f"    ✅ Found responsibilities heading!")
                    elif any(bullet in text for bullet in ['•', '●', '◦', '*', '-']):
                        print(f"    📋 Bullet point detected")
                    elif len(text.split()) <= 6 and any(keyword in text.lower() for keyword in ['development', 'project', 'system', 'application']):
                        print(f"    🎯 Potential project name")
                        
        except Exception as e:
            print(f"❌ Error reading {resume_file}: {e}")
    
    # Now test the project detector
    print(f"\n🔍 Testing ProjectDetector on files")
    print("=" * 50)
    
    try:
        from resume_customizer.detectors.project_detector import ProjectDetector
        
        detector = ProjectDetector()
        
        for resume_file in resume_files:
            if not os.path.exists(resume_file):
                continue
                
            print(f"\n📄 ProjectDetector analysis of {resume_file}")
            print("-" * 40)
            
            doc = Document(resume_file)
            projects = detector.find_projects(doc)
            
            print(f"Found {len(projects)} projects:")
            for j, project in enumerate(projects):
                print(f"  Project {j+1}:")
                print(f"    Name: {project.name}")
                print(f"    Role: {project.role}")
                print(f"    Company: {project.company}")
                print(f"    Date Range: {project.date_range}")
                print(f"    Bullet Points: {len(project.bullet_points)}")
                for bp in project.bullet_points[:3]:  # Show first 3
                    print(f"      • {bp}")
                if len(project.bullet_points) > 3:
                    print(f"      ... and {len(project.bullet_points) - 3} more")
                print()
                
    except Exception as e:
        print(f"❌ Error testing ProjectDetector: {e}")
        import traceback
        traceback.print_exc()
        
    print("=" * 50)
    print("🎯 SUMMARY")
    print("If no projects with Responsibilities sections were found,")
    print("the issue is likely in the document structure or detector logic.")
    print("=" * 50)

if __name__ == "__main__":
    debug_resume_structure()