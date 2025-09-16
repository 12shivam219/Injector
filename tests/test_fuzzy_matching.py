"""
Tests for the fuzzy company name matching functionality
"""

import os
import pytest
from docx import Document
from docx.shared import Pt

from resume_customizer.analyzers.format_analyzer import FormatAnalyzer
from utilities.fuzzy_matcher import FuzzyMatcher
from database.format_models import ResumeFormat

def create_test_doc(company_names: list) -> Document:
    """Create a test document with given company names"""
    doc = Document()
    
    # Add experience section
    doc.add_heading("EXPERIENCE", 1)
    
    # Add companies with some formatting
    for company in company_names:
        p = doc.add_paragraph()
        run = p.add_run(company)
        run.bold = True
        doc.add_paragraph("Some role")  # Add dummy role
        doc.add_paragraph("• Some accomplishment")
    
    return doc

def test_fuzzy_matcher_basic():
    """Test basic fuzzy matching functionality"""
    matcher = FuzzyMatcher()
    
    # Test exact match
    assert matcher.compare_strings("Boeing", "Boeing")[0] == True
    
    # Test company suffixes
    is_match, score = matcher.compare_strings("Boeing", "Boeing Inc.")
    assert is_match == True
    assert score >= 85
    
    is_match, score = matcher.compare_strings("Boeing", "The Boeing Company")
    assert is_match == True
    assert score >= 85
    
    # Test minor variations
    is_match, score = matcher.compare_strings("Mastery Logistics", "Mastery Logistics LLC")
    assert is_match == True
    assert score >= 85
    
    # Test different word order
    is_match, score = matcher.compare_strings("Logistics Mastery", "Mastery Logistics")
    assert is_match == True
    assert score >= 85
    
    # Test case sensitivity
    is_match, score = matcher.compare_strings("BOEING", "boeing")
    assert is_match == True
    assert score >= 85

def test_fuzzy_matcher_negative_cases():
    """Test cases that should not match"""
    matcher = FuzzyMatcher()
    
    # Test completely different names
    assert matcher.compare_strings("Boeing", "Airbus")[0] == False
    
    # Test similar but different enough names
    is_match, score = matcher.compare_strings("Zapier", "Zapper")
    assert is_match == False
    assert score < 85
    
    # Test substring that shouldn't match
    assert matcher.compare_strings("Tech", "TechCorp")[0] == False

def test_format_analyzer_company_matching():
    """Test company matching in format analyzer"""
    analyzer = FormatAnalyzer()
    
    # Create test documents
    original_companies = ["Boeing Corporation", "Mastery Logistics LLC", "Cardinal Health"]
    original_doc = create_test_doc(original_companies)
    
    variant_companies = ["The Boeing Company", "Mastery Logistics", "CardinalHealthcare"]
    variant_doc = create_test_doc(variant_companies)
    
    # Create a mock stored format
    stored_format = ResumeFormat()
    stored_format.company_patterns = [{
        'pattern': name,
        'formatting': {'bold': True}
    } for name in original_companies]
    
    # Extract patterns from variant document
    new_patterns = analyzer.analyze_format(variant_doc)
    
    # Test matching
    match_score, matched_elements = analyzer.match_format(variant_doc, stored_format)
    
    # We expect a high company match score due to fuzzy matching
    assert matched_elements['companies'] >= 0.8
    
    # The overall match score should be good
    assert match_score >= 70

def test_fuzzy_matcher_edge_cases():
    """Test edge cases for fuzzy matching"""
    matcher = FuzzyMatcher()
    
    # Test empty strings
    assert matcher.compare_strings("", "")[0] == False
    assert matcher.compare_strings("Boeing", "")[0] == False
    assert matcher.compare_strings("", "Boeing")[0] == False
    
    # Test strings with only spaces
    assert matcher.compare_strings("   ", "")[0] == False
    assert matcher.compare_strings("Boeing", "   ")[0] == False
    
    # Test special characters
    is_match, score = matcher.compare_strings("Company-Name", "CompanyName")
    assert is_match == True
    assert score >= 85
    
    # Test very long company names
    long_name1 = "International Business Machines Corporation Global Services Division"
    long_name2 = "IBM Global Services Division"
    is_match, score = matcher.compare_strings(long_name1, long_name2)
    assert is_match == True
    assert score >= 85

if __name__ == "__main__":
    pytest.main([__file__])