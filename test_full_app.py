"""
Full application test script to test the complete workflow with user's actual resume
"""

import sys
import os
sys.path.append(os.getcwd())

from io import BytesIO
from resume_customizer.processors.resume_processor import ResumeManager

def test_full_application_workflow():
    """Test the complete application workflow with user's actual resume"""
    print("=== FULL APPLICATION WORKFLOW TEST ===\n")
    
    resume_path = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    # Check if file exists
    if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return False
    
    print(f"📄 Testing with: {os.path.basename(resume_path)}")
    
    try:
        # Step 1: Load resume file
        print("\n1. Loading resume file...")
        with open(resume_path, 'rb') as file:
            file_buffer = BytesIO(file.read())
        print("✅ Resume loaded successfully")
        
        # Step 2: Prepare tech stack input (Java-focused for your resume)
        print("\n2. Preparing tech stack input...")
        tech_input = """Java
• Developed enterprise-grade applications using Java 8-21 features
• Implemented microservices architecture with Spring Boot
• Created RESTful APIs with comprehensive error handling
• Applied design patterns for scalable application development

Spring Framework
• Built microservices using Spring Boot with auto-configuration
• Implemented Spring Security for authentication and authorization
• Developed batch processing solutions with Spring Batch
• Used Spring MVC for web application development

AWS
• Deployed applications on EC2 instances with auto-scaling
• Utilized Lambda functions for serverless computing
• Configured RDS for database management and backups
• Implemented CloudWatch for monitoring and logging

Docker
• Containerized Java applications for consistent deployments
• Created multi-stage Docker builds for optimization
• Orchestrated containers using Docker Compose
• Integrated Docker with CI/CD pipelines"""

        print("✅ Tech stack prepared with Java, Spring Framework, AWS, Docker")
        
        # Step 3: Generate preview using ResumeManager
        print("\n3. Generating resume preview...")
        manager = ResumeManager()
        
        preview_result = manager.generate_preview(
            file_obj=file_buffer,
            input_text=tech_input,
            manual_text=""
        )
        
        if not preview_result['success']:
            print(f"❌ Preview generation failed: {preview_result['error']}")
            if 'debug_info' in preview_result:
                print(f"🐛 Debug info: {preview_result['debug_info']}")
            return False
        
        print("✅ Preview generated successfully!")
        print(f"📊 Points added: {preview_result['points_added']}")
        print(f"📊 Projects found: {preview_result['projects_count']}")
        print(f"📊 Tech stacks used: {', '.join(preview_result['tech_stacks_used'])}")
        
        # Display point distribution
        print(f"\n🎯 Points Distribution:")
        for project, mapping in preview_result['project_points_mapping'].items():
            print(f"  📋 {project}: {len(mapping['points'])} points")
            for i, point in enumerate(mapping['points'], 1):
                print(f"     {i}. {point}")
        
        # Step 4: Process full resume (simulating the actual processing)
        print(f"\n4. Processing full resume...")
        
        # Reset buffer for processing
        file_buffer.seek(0)
        
        # Prepare file data as the app would
        file_data = {
            'filename': 'Resume Format 1.docx',
            'file': file_buffer,
            'text': tech_input,
            'manual_text': '',
            'recipient_email': '',
            'sender_email': '',
            'sender_password': '',
            'smtp_server': '',
            'smtp_port': 465,
            'email_subject': '',
            'email_body': ''
        }
        
        # Process the resume
        process_result = manager.process_single_resume(file_data)
        
        if not process_result['success']:
            print(f"❌ Resume processing failed: {process_result['error']}")
            return False
        
        print("✅ Resume processed successfully!")
        print(f"📊 Points added to final resume: {process_result['points_added']}")
        print(f"📊 Tech stacks applied: {', '.join(process_result['tech_stacks'])}")
        
        # Step 5: Save the processed resume
        print(f"\n5. Saving processed resume...")
        output_path = r"C:\Users\HP\Downloads\RSInjector\processed_resume_output.docx"
        
        if 'buffer' in process_result:
            with open(output_path, 'wb') as output_file:
                output_file.write(process_result['buffer'])
            print(f"💾 Processed resume saved to: {output_path}")
        else:
            print("❌ No processed content to save")
            return False
        
        print(f"\n✅ FULL WORKFLOW COMPLETED SUCCESSFULLY!")
        print(f"📈 Summary:")
        print(f"   • Projects detected: {preview_result['projects_count']}")
        print(f"   • Tech stacks processed: {len(preview_result['tech_stacks_used'])}")
        print(f"   • Points added: {process_result['points_added']}")
        print(f"   • Output file: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error in workflow: {e}")
        import traceback
        traceback.print_exc()
        return False

def compare_before_after():
    """Compare the original and processed resumes"""
    print("\n=== BEFORE/AFTER COMPARISON ===\n")
    
    original_path = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    processed_path = r"C:\Users\HP\Downloads\RSInjector\processed_resume_output.docx"
    
    try:
        from docx import Document
        
        print("📄 Original Resume Structure:")
        original_doc = Document(original_path)
        original_paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
        print(f"   Total paragraphs: {len(original_paragraphs)}")
        
        if os.path.exists(processed_path):
            print(f"\n📄 Processed Resume Structure:")
            processed_doc = Document(processed_path)
            processed_paragraphs = [p.text.strip() for p in processed_doc.paragraphs if p.text.strip()]
            print(f"   Total paragraphs: {len(processed_paragraphs)}")
            
            # Count new bullet points
            original_bullets = [p for p in original_paragraphs if p.startswith('-')]
            processed_bullets = [p for p in processed_paragraphs if p.startswith('-')]
            
            print(f"\n📊 Bullet Point Comparison:")
            print(f"   Original bullets: {len(original_bullets)}")
            print(f"   Processed bullets: {len(processed_bullets)}")
            print(f"   New bullets added: {len(processed_bullets) - len(original_bullets)}")
            
            # Show some new bullets if any
            if len(processed_bullets) > len(original_bullets):
                print(f"\n🆕 Some newly added bullets:")
                new_bullets = processed_bullets[len(original_bullets):]
                for i, bullet in enumerate(new_bullets[:5], 1):
                    print(f"   {i}. {bullet}")
                if len(new_bullets) > 5:
                    print(f"   ... and {len(new_bullets) - 5} more")
        else:
            print("❌ Processed resume not found")
            
    except Exception as e:
        print(f"❌ Error comparing files: {e}")

if __name__ == "__main__":
    # Run the full workflow test
    success = test_full_application_workflow()
    
    if success:
        # Compare before and after
        compare_before_after()
        
        print(f"\n🎉 ALL TESTS PASSED!")
        print(f"🚀 Your resume is fully compatible with the RSInjector application!")
        print(f"💡 You can now use the web interface at: http://localhost:5000")
        print(f"📝 Simply upload your resume and add your tech stacks to get enhanced results!")
    else:
        print(f"\n❌ Some tests failed. Check the output above for details.")