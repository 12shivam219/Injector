"""
Comprehensive test script for Resume Format 1 and Resume Format 3
Tests project detection, point injection, and overall functionality
"""

import sys
import os
sys.path.append(os.getcwd())

from io import BytesIO
from docx import Document
from resume_customizer.processors.resume_processor import ResumeManager
from resume_customizer.detectors.project_detector import ProjectDetector

def analyze_resume_structure(file_path, resume_name):
    """Analyze and display resume structure"""
    print(f"\n{'='*60}")
    print(f"📄 ANALYZING {resume_name}")
    print(f"{'='*60}")
    
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        print(f"📊 Total paragraphs: {len(paragraphs)}")
        
        # Find potential project sections
        project_indicators = []
        responsibility_indicators = []
        
        for i, para in enumerate(paragraphs):
            if any(keyword in para.lower() for keyword in ['project', 'experience', 'work', 'position']):
                if len(para) < 100:  # Likely a header
                    project_indicators.append((i, para))
            
            if 'responsibilit' in para.lower():
                responsibility_indicators.append((i, para))
        
        print(f"🎯 Potential project headers: {len(project_indicators)}")
        for i, (idx, text) in enumerate(project_indicators[:5]):  # Show first 5
            print(f"   {i+1}. Line {idx}: {text[:80]}...")
        
        print(f"📝 Responsibility sections: {len(responsibility_indicators)}")
        for i, (idx, text) in enumerate(responsibility_indicators):
            print(f"   {i+1}. Line {idx}: {text[:60]}...")
        
        # Count bullet points
        bullets = [p for p in paragraphs if p.startswith(('-', '•', '*', '◦', '▪'))]
        print(f"🔸 Bullet points found: {len(bullets)}")
        
        return doc, paragraphs
        
    except Exception as e:
        print(f"❌ Error analyzing {resume_name}: {e}")
        return None, None

def test_project_detection(file_path, resume_name):
    """Test project detection functionality"""
    print(f"\n🔍 TESTING PROJECT DETECTION - {resume_name}")
    print("-" * 50)
    
    try:
        with open(file_path, 'rb') as file:
            file_buffer = BytesIO(file.read())
        
        doc = Document(file_buffer)
        detector = ProjectDetector()
        projects = detector.find_projects(doc)
        
        print(f"✅ Projects detected: {len(projects)}")
        
        for i, project in enumerate(projects, 1):
            print(f"\n📋 Project {i}:")
            print(f"   Name: {project.name}")
            print(f"   Company: {project.company}")
            print(f"   Role: {project.role}")
            print(f"   Date Range: {project.date_range}")
            print(f"   Start Index: {project.start_index}")
            print(f"   End Index: {project.end_index}")
            print(f"   Bullet Points: {len(project.bullet_points)}")
            
            # Show some bullet points
            if project.bullet_points:
                print("   Sample bullets:")
                for j, bullet in enumerate(project.bullet_points[:3]):
                    print(f"     • {bullet[:60]}...")
        
        return projects
        
    except Exception as e:
        print(f"❌ Project detection failed: {e}")
        import traceback
        traceback.print_exc()
        return []

def test_full_processing(file_path, resume_name):
    """Test full resume processing with tech stack injection"""
    print(f"\n🚀 TESTING FULL PROCESSING - {resume_name}")
    print("-" * 50)
    
    try:
        # Prepare tech stack input
        tech_input = """Java
• Developed enterprise-grade applications using Java 8-21 features
• Implemented microservices architecture with Spring Boot
• Created RESTful APIs with comprehensive error handling
• Applied design patterns for scalable application development

Spring Framework
• Built microservices using Spring Boot with auto-configuration
• Implemented Spring Security for authentication and authorization
• Developed batch processing solutions with Spring Batch
• Used Spring MVC for web application development

AWS
• Deployed applications on EC2 instances with auto-scaling
• Utilized Lambda functions for serverless computing
• Configured RDS for database management and backups
• Implemented CloudWatch for monitoring and logging

React
• Developed responsive single-page applications
• Implemented component-based architecture patterns
• Utilized Redux for state management
• Created reusable UI component libraries"""

        print("📝 Tech stack input prepared (Java, Spring, AWS, React)")
        
        # Load file
        with open(file_path, 'rb') as file:
            file_buffer = BytesIO(file.read())
        
        # Test preview generation
        manager = ResumeManager()
        preview_result = manager.generate_preview(
            file_obj=file_buffer,
            input_text=tech_input,
            manual_text=""
        )
        
        if not preview_result['success']:
            print(f"❌ Preview generation failed: {preview_result['error']}")
            return None
        
        print("✅ Preview generated successfully!")
        print(f"📊 Points added: {preview_result['points_added']}")
        print(f"📊 Projects count: {preview_result['projects_count']}")
        print(f"📊 Tech stacks used: {', '.join(preview_result['tech_stacks_used'])}")
        
        # Show point distribution
        print(f"\n🎯 Points Distribution:")
        for project, mapping in preview_result['project_points_mapping'].items():
            print(f"📋 {project}: {len(mapping['points'])} points")
            for i, point in enumerate(mapping['points'], 1):
                print(f"     {i}. {point}")
        
        # Test actual processing
        print(f"\n🔄 Processing full resume...")
        file_buffer.seek(0)  # Reset buffer
        
        file_data = {
            'filename': os.path.basename(file_path),
            'file': file_buffer,
            'text': tech_input,
            'manual_text': '',
            'recipient_email': '',
            'sender_email': '',
            'sender_password': '',
            'smtp_server': '',
            'smtp_port': 465,
            'email_subject': '',
            'email_body': ''
        }
        
        process_result = manager.process_single_resume(file_data)
        
        if not process_result['success']:
            print(f"❌ Resume processing failed: {process_result['error']}")
            return None
        
        print("✅ Resume processed successfully!")
        print(f"📊 Final points added: {process_result['points_added']}")
        print(f"📊 Tech stacks applied: {', '.join(process_result['tech_stacks'])}")
        
        # Save processed resume
        output_filename = f"processed_{os.path.basename(file_path)}"
        output_path = os.path.join(os.getcwd(), output_filename)
        
        if 'buffer' in process_result:
            with open(output_path, 'wb') as output_file:
                output_file.write(process_result['buffer'])
            print(f"💾 Processed resume saved to: {output_path}")
            
            # Analyze the processed resume
            analyze_processed_resume(output_path, resume_name)
            
        return process_result
        
    except Exception as e:
        print(f"❌ Full processing failed: {e}")
        import traceback
        traceback.print_exc()
        return None

def analyze_processed_resume(processed_path, resume_name):
    """Analyze the processed resume to verify changes"""
    print(f"\n📊 ANALYZING PROCESSED RESUME - {resume_name}")
    print("-" * 50)
    
    try:
        doc = Document(processed_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Count bullet points
        bullets = [p for p in paragraphs if p.startswith(('-', '•', '*', '◦', '▪'))]
        
        # Look for newly added tech-specific bullet points
        tech_bullets = []
        tech_keywords = ['Java', 'Spring', 'AWS', 'React', 'microservices', 'REST', 'API', 'Lambda', 'EC2']
        
        for bullet in bullets:
            if any(keyword in bullet for keyword in tech_keywords):
                tech_bullets.append(bullet)
        
        print(f"📄 Total paragraphs in processed resume: {len(paragraphs)}")
        print(f"🔸 Total bullet points: {len(bullets)}")
        print(f"💻 Tech-specific bullets found: {len(tech_bullets)}")
        
        if tech_bullets:
            print("\n🆕 Sample new tech bullets added:")
            for i, bullet in enumerate(tech_bullets[:5], 1):
                print(f"   {i}. {bullet}")
        
        return len(tech_bullets)
        
    except Exception as e:
        print(f"❌ Error analyzing processed resume: {e}")
        return 0

def main():
    """Main test function"""
    print("🧪 COMPREHENSIVE RESUME TESTING")
    print("="*60)
    
    resume_files = [
        ("C:\\Users\\HP\\Downloads\\Resume Format 1.docx", "Resume Format 1"),
        ("C:\\Users\\HP\\Downloads\\Resume Format 3.DOCX", "Resume Format 3")
    ]
    
    results = {}
    
    for file_path, resume_name in resume_files:
        print(f"\n{'🎯 TESTING ' + resume_name + ' ':=^60}")
        
        # Step 1: Analyze structure
        doc, paragraphs = analyze_resume_structure(file_path, resume_name)
        if not doc:
            continue
        
        # Step 2: Test project detection
        projects = test_project_detection(file_path, resume_name)
        
        # Step 3: Test full processing
        process_result = test_full_processing(file_path, resume_name)
        
        # Store results
        results[resume_name] = {
            'projects_detected': len(projects),
            'processing_success': process_result is not None,
            'points_added': process_result.get('points_added', 0) if process_result else 0
        }
    
    # Final summary
    print(f"\n{'🎉 FINAL SUMMARY ':=^60}")
    for resume_name, result in results.items():
        print(f"\n📄 {resume_name}:")
        print(f"   Projects detected: {result['projects_detected']}")
        print(f"   Processing successful: {'✅ Yes' if result['processing_success'] else '❌ No'}")
        print(f"   Points added: {result['points_added']}")
        print(f"   Status: {'🟢 PASSED' if result['processing_success'] and result['projects_detected'] > 0 else '🔴 FAILED'}")
    
    print(f"\n{'TEST COMPLETED':=^60}")

if __name__ == "__main__":
    main()