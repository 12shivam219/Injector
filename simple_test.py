import docx
from resume_customizer.detectors.project_detector import ProjectDetector
from resume_customizer.processors.document_processor import DocumentProcessor

def main():
    print("Testing project detection and point placement fixes...")
    
    # Create a test document
    doc = docx.Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("ABC Company | Jan 2020 - Present")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("• Developed web applications")
    doc.add_paragraph("• Implemented CI/CD pipelines")
    doc.add_paragraph("XYZ Corporation | Mar 2018 - Dec 2019")
    doc.add_paragraph("Software Developer")
    doc.add_paragraph("• Built RESTful APIs")
    doc.add_paragraph("• Collaborated with teams")
    
    # Save the test document
    test_file = "test_resume.docx"
    doc.save(test_file)
    print(f"Created test document: {test_file}")
    
    # Test project detection
    detector = ProjectDetector()
    projects = detector.find_projects(doc)
    
    print(f"\nDetected {len(projects)} projects:")
    for i, project in enumerate(projects):
        print(f"Project {i+1}:")
        print(f"  Name: {project.name}")
        print(f"  Role: {project.role}")
        print(f"  Start Index: {project.start_index}")
        print(f"  End Index: {project.end_index}")
        print(f"  Bullet Points: {project.bullet_points}")
    
    # Test point placement
    processor = DocumentProcessor()
    new_doc = docx.Document(test_file)
    
    # Use the actual ProjectInfo object instead of a dict
    project = projects[0]
    
    # Add points to the first project
    new_points = [{"text": "• Added new React components for improved UX"}]
    document_marker = "•"  # Default bullet marker
    processor._add_points_to_project(new_doc, project, new_points, document_marker)
    
    # Save the modified document
    modified_file = "test_resume_modified.docx"
    new_doc.save(modified_file)
    print(f"\nAdded points to document: {modified_file}")
    
    # Print the modified document content
    print("\nModified document content:")
    for i, para in enumerate(new_doc.paragraphs):
        print(f"{i}: {para.text}")
    
    print("\nTest completed. Please check the modified document to verify point placement.")

if __name__ == "__main__":
    main()