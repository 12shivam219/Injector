"""
Format Management Tab for Resume Customizer
Handles resume format uploading and management
"""

import streamlit as st
from typing import Dict, Any, Optional
import docx
from datetime import datetime

from resume_customizer.analyzers.format_analyzer import FormatAnalyzer
from database.format_models import ResumeFormat, ResumeFormatMatch
from infrastructure.storage.local_format_store import LocalFormatStore
from infrastructure.utilities.logger import get_logger

logger = get_logger()

class FormatManager:
    """Manages resume format templates and matching"""
    
    def __init__(self, session_factory):
        """Initialize Format Manager with session factory"""
        self.session_maker = session_factory
        self.analyzer = FormatAnalyzer()
        # Local fallback store
        self.local_store = LocalFormatStore()
        
    def render_format_tab(self):
        """Render the format management tab"""
        st.title("📋 Resume Format Manager")
        st.markdown("### Upload and manage resume format templates")
        
        # Create tabs for different functions
        tab1, tab2, tab3 = st.tabs([
            "📤 Upload New Format",
            "📊 View Formats",
            "🔍 Test Format Matching"
        ])
        
        with tab1:
            self._render_upload_section()
            
        with tab2:
            self._render_formats_list()
            # Add migration option when local formats exist but DB is available
            if self.session_maker:
                try:
                    local_items = self.local_store.list_formats()
                    if local_items:
                            st.markdown("**Local templates detected.** You can migrate them into the database.")
                            remove_after = st.checkbox("Remove local templates after successful migration?", value=False)
                            if st.button("⬆️ Migrate local templates to database"):
                                # Ask for confirmation
                                if st.confirm("Are you sure you want to migrate local templates to the database?"):
                                    self._migrate_local_to_db(local_items, remove_after=remove_after)
                except Exception:
                    pass
            
        with tab3:
            self._render_test_matching()
    
    def _render_upload_section(self):
        """Render the format upload section"""
        st.markdown("### Upload a New Format Template")
        st.info("Upload a sample resume to extract its format patterns")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a resume file (DOCX)",
            type=["docx"],
            key="format_uploader"
        )
        
        # Format name and description
        format_name = st.text_input("Format Name", 
            placeholder="e.g., Standard Professional Format"
        )
        
        format_desc = st.text_area("Format Description",
            placeholder="Describe the format's characteristics..."
        )
        
        if uploaded_file and format_name:
            if st.button("📥 Analyze and Save Format"):
                # Process the upload
                doc = docx.Document(uploaded_file)
                patterns = self.analyzer.analyze_format(doc)
                
                # Create format record
                new_format = ResumeFormat(
                    name=format_name,
                    description=format_desc,
                    name_pattern=patterns.get('name_pattern', {}),
                    email_pattern=patterns.get('email_pattern', {}),
                    phone_pattern=patterns.get('phone_pattern', {}),
                    section_patterns=patterns.get('section_patterns', {}),
                    company_patterns=patterns.get('company_patterns', []),
                    title_patterns=patterns.get('title_patterns', []),
                    last_used=datetime.now()
                )
                
                # Save to database using session if available, otherwise local store
                if self.session_maker:
                    session = self.session_maker()
                    try:
                        session.add(new_format)
                        session.commit()
                        st.success(f"✅ Format '{format_name}' has been saved!")
                    except Exception as e:
                        st.error(f"❌ Error saving format: {str(e)}")
                        logger.error(f"Error in format upload: {str(e)}")
                    finally:
                        session.close()
                else:
                    try:
                        fid = self.local_store.save_format({
                            'name': format_name,
                            'description': format_desc,
                            'name_pattern': patterns.get('name_pattern', {}),
                            'email_pattern': patterns.get('email_pattern', {}),
                            'phone_pattern': patterns.get('phone_pattern', {}),
                            'section_patterns': patterns.get('section_patterns', {}),
                            'company_patterns': patterns.get('company_patterns', []),
                            'title_patterns': patterns.get('title_patterns', []),
                            'last_used': datetime.now().isoformat()
                        })
                        st.success(f"✅ Format '{format_name}' saved locally (id: {fid})")
                    except Exception as e:
                        st.error(f"❌ Error saving local format: {str(e)}")
                        logger.error(f"Error saving local format: {str(e)}")
    
    def _render_formats_list(self):
        """Render the list of stored formats"""
        st.markdown("### Stored Format Templates")
        
        # If DB session is available, use DB, otherwise use local fallback
        if self.session_maker:
            session = self.session_maker()
            try:
                formats = session.query(ResumeFormat).all()

                if not formats:
                    st.warning("No format templates stored yet")
                    return

                # Display formats
                for fmt in formats:
                    with st.expander(f"📄 {fmt.name}"):
                        col1, col2 = st.columns([3,1])

                        with col1:
                            st.markdown(f"**Description:** {fmt.description}")
                            st.markdown(f"**Match Count:** {fmt.match_count} resumes")
                            st.markdown(f"**Last Used:** {fmt.last_used}")

                        with col2:
                            if st.button("🗑️ Delete", key=f"delete_{fmt.id}"):
                                try:
                                    session.delete(fmt)
                                    session.commit()
                                    st.success("Format deleted!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error deleting format: {str(e)}")
                        # Show pattern details
                        if st.checkbox("Show Pattern Details", key=f"details_{fmt.id}"):
                            st.json({
                                'sections': fmt.section_patterns,
                                'companies': fmt.company_patterns,
                                'titles': fmt.title_patterns
                            })
            finally:
                session.close()
        else:
            # Local fallback
            formats = self.local_store.list_formats()
            if not formats:
                st.warning("No format templates stored yet (local fallback)")
                return

            for fmt in formats:
                with st.expander(f"📄 {fmt.get('name')}"):
                    col1, col2 = st.columns([3,1])

                    with col1:
                        st.markdown(f"**Description:** {fmt.get('description')}")
                        st.markdown(f"**Match Count:** {fmt.get('match_count')} resumes")
                        st.markdown(f"**Last Used:** {fmt.get('last_used')}")

                    with col2:
                        if st.button("🗑️ Delete", key=f"delete_local_{fmt.get('id')}"):
                            try:
                                if self.local_store.delete_format(fmt.get('id')):
                                    st.success("Format deleted (local)!")
                                    st.experimental_rerun()
                                else:
                                    st.error("Could not delete local format")
                            except Exception as e:
                                st.error(f"Error deleting local format: {str(e)}")
                    if st.checkbox("Show Pattern Details", key=f"details_local_{fmt.get('id')}"):
                        st.json({
                            'sections': fmt.get('section_patterns'),
                            'companies': fmt.get('company_patterns'),
                            'titles': fmt.get('title_patterns')
                        })
    
    def _render_test_matching(self):
        """Render the format matching test section"""
        st.markdown("### Test Format Matching")
        st.info("Upload a resume to test which format it matches")
        
        # File uploader for test
        test_file = st.file_uploader(
            "Choose a resume to test",
            type=["docx"],
            key="test_uploader"
        )
        
        # Use DB session when available, otherwise local store for formats
        if self.session_maker:
            session = self.session_maker()
            try:
                if test_file:
                    doc = docx.Document(test_file)
                    formats = session.query(ResumeFormat).all()

                    if not formats:
                        st.warning("No formats available for matching")
                        return

                    # Test against all formats
                    results = []
                    for fmt in formats:
                        score, elements = self.analyzer.match_format(doc, fmt)
                        results.append({
                            'format': fmt.name,
                            'score': score,
                            'elements': elements
                        })

                    # Sort by score
                    results.sort(key=lambda x: x['score'], reverse=True)

                    # Show results
                    st.markdown("### Match Results")
                    for result in results:
                        score_color = 'green' if result['score'] >= 70 else 'orange' if result['score'] >= 40 else 'red'
                        st.markdown(
                            f"**{result['format']}**: "
                            f"<span style='color: {score_color}'>{result['score']}%</span> match",
                            unsafe_allow_html=True
                        )

                        if st.checkbox(f"Show details for {result['format']}", key=f"result_{result['format']}"):
                            st.json(result['elements'])
            finally:
                session.close()
        
        else:
            if test_file:
                doc = docx.Document(test_file)
                formats = self.local_store.list_formats()

                if not formats:
                    st.warning("No formats available for matching (local fallback)")
                    return

                results = []
                for fmt in formats:
                    # local format dict -> adapt to analyzer expectations
                    score, elements = self.analyzer.match_format(doc, fmt)
                    results.append({
                        'format': fmt.get('name'),
                        'score': score,
                        'elements': elements
                    })

                results.sort(key=lambda x: x['score'], reverse=True)

                st.markdown("### Match Results")
                for result in results:
                    score_color = 'green' if result['score'] >= 70 else 'orange' if result['score'] >= 40 else 'red'
                    st.markdown(
                        f"**{result['format']}**: "
                        f"<span style='color: {score_color}'>{result['score']}%</span> match",
                        unsafe_allow_html=True
                    )

                    if st.checkbox(f"Show details for {result['format']}", key=f"result_{result['format']}"):
                        st.json(result['elements'])
    
    def _migrate_local_to_db(self, local_items, remove_after: bool = False):
        """Migrate formats stored locally into the database"""
        if not self.session_maker:
            st.error("Database session not available for migration")
            return

        session = self.session_maker()
        migrated = 0
        try:
            for item in local_items:
                # Check for existing by name
                exists = session.query(ResumeFormat).filter_by(name=item.get('name')).first()
                if exists:
                    continue

                rf = ResumeFormat(
                    name=item.get('name'),
                    description=item.get('description'),
                    name_pattern=item.get('name_pattern'),
                    email_pattern=item.get('email_pattern'),
                    phone_pattern=item.get('phone_pattern'),
                    section_patterns=item.get('section_patterns'),
                    company_patterns=item.get('company_patterns'),
                    title_patterns=item.get('title_patterns'),
                    last_used=item.get('last_used')
                )
                session.add(rf)
                migrated += 1
            session.commit()
            st.success(f"✅ Migrated {migrated} local templates to database")
            if remove_after and migrated > 0:
                for item in local_items:
                    try:
                        self.local_store.delete_format(item.get('id'))
                    except Exception:
                        pass
        except Exception as e:
            session.rollback()
            st.error(f"Migration failed: {e}")
        finally:
            session.close()

def main():
    """Main entry point for the Format Manager page"""
    # Initialize the app bootstrap if needed
    from infrastructure.app.app_bootstrap import initialize_app, get_cached_services
    
    initialize_app()
    
    # Get services including database session
    services = get_cached_services()
    db_session = services.get('db_session')
    
    if not db_session:
        st.error("❌ Database connection not available")
        return
        
    # Create and render format manager
    manager = FormatManager(db_session)
    manager.render_format_tab()

if __name__ == "__main__":
    main()